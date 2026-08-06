#!/usr/bin/env python3
"""Genera documento .docx con la especificación de módulos de programación de VERIFEX."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.dirname(os.path.abspath(__file__))

# Paleta de colores profesional
NAVY = "1B3A5C"
DARK_SLATE = "2C3E50"
MEDIUM_SLATE = "34495E"
DARK_RED = "8B0000"
HEADING_COLOR = RGBColor(0x1B, 0x3A, 0x5C)
SUBTITLE_GRAY = RGBColor(0x70, 0x70, 0x80)
NOTE_COLOR = RGBColor(0xC0, 0x39, 0x2B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_attr_row(table, *cell_texts):
    row = table.add_row()
    for i, text in enumerate(cell_texts):
        if i < len(row.cells):
            row.cells[i].text = text
    for cell in row.cells:
        for p in cell.paragraphs:
            p.style.font.size = Pt(8.5)
            p.style.font.name = "Consolas"


def create_table(doc, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = 1
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = "Consolas"
            run.font.color.rgb = WHITE
        set_cell_shading(cell, DARK_SLATE)
    return table


def add_title_page(doc):
    for _ in range(6):
        doc.add_paragraph("")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("VERIFEX")
    run.font.size = Pt(42)
    run.bold = True
    run.font.color.rgb = HEADING_COLOR

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Especificación de Módulos de Programación")
    run.font.size = Pt(18)
    run.font.color.rgb = SUBTITLE_GRAY

    doc.add_paragraph("")
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub2.add_run("Clases, Métodos, Funciones y Componentes")
    run.font.size = Pt(14)
    run.font.color.rgb = SUBTITLE_GRAY

    doc.add_paragraph("")
    sub3 = doc.add_paragraph()
    sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub3.add_run("Backend Python/Flask + Frontend React/TypeScript")
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = SUBTITLE_GRAY

    doc.add_page_break()


def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = HEADING_COLOR
    return h


def add_backend_section(doc):
    add_heading_styled(doc, "1. Backend (Python/Flask)", 0)
    add_heading_styled(doc, "1.1 app.py — Módulo de la Aplicación Flask", 1)

    p = doc.add_paragraph()
    run = p.add_run("Archivo: ")
    run.bold = True
    p.add_run("server/app.py").font.name = "Consolas"

    p = doc.add_paragraph(
        "Punto de entrada del servidor Flask. Define las rutas HTTP, "
        "configura CORS, carga variables de entorno y sirve el frontend "
        "compilado como SPA (Single Page Application)."
    )
    doc.add_paragraph("")

    add_heading_styled(doc, "Variables del Módulo", 2)
    tbl = create_table(doc, ["Visibilidad", "Nombre", "Tipo", "Descripción"])
    add_attr_row(tbl, "pública", "app", "Flask",
                 "Instancia principal de Flask. Configurada con static_folder='../dist' y CORS habilitado para todos los orígenes.")
    doc.add_paragraph("")

    add_heading_styled(doc, "Rutas (Endpoints HTTP)", 2)
    tbl = create_table(doc, ["Método", "Ruta", "Función", "Parámetros", "Retorno", "Descripción"])
    add_attr_row(tbl, "POST", "/analyze", "analyze()", "— (lee JSON del body: {url: string})", "flask.Response (JSON)",
                 "Valida la URL, llama a analyze_url() del módulo analyzer, llama a find_similar_news() para noticias relacionadas, y retorna el análisis completo como JSON.")
    add_attr_row(tbl, "GET", "/health", "health()", "—", "flask.Response (JSON)",
                 "Endpoint de verificación de salud. Retorna {'status': 'ok'}. Usado por monitoreo y health checks de Render.")
    add_attr_row(tbl, "GET", "/<path:path>", "serve_frontend(path)", "path: str — ruta solicitada", "flask.Response (archivo estático o index.html)",
                 "Sirve el frontend React compilado (SPA). Si el archivo existe en dist/, lo sirve directamente. Caso contrario, sirve index.html para que React Router maneje la ruta.")
    doc.add_paragraph("")

    add_heading_styled(doc, "1.2 analyzer.py — Módulo de Análisis de Credibilidad", 1)

    p = doc.add_paragraph()
    run = p.add_run("Archivo: ")
    run.bold = True
    p.add_run("server/analyzer.py").font.name = "Consolas"
    p = doc.add_paragraph(
        "Módulo principal del sistema. Contiene toda la lógica de extracción "
        "de contenido web (scraping multicapa), comunicación con la API de Groq "
        "(LLM) para clasificación, y construcción de la respuesta final."
    )
    doc.add_paragraph("")

    add_heading_styled(doc, "1.2.1 Constantes del Módulo", 2)
    tbl = create_table(doc, ["Nombre", "Tipo", "Descripción"])
    add_attr_row(tbl, "CREDIBLE_DOMAINS", "Set[str]",
                 "Conjunto de dominios de medios de comunicación reconocidos (jornada.com.mx, milenio.com, reuters.com, etc.). Si el dominio de la URL está en este conjunto y el veredicto es FALSO, se reclasifica automáticamente a NO VERIFICABLE.")
    add_attr_row(tbl, "SOCIAL_MEDIA_DOMAINS", "Set[str]",
                 "Dominios de redes sociales (instagram.com, x.com, facebook.com, tiktok.com, etc.). Reciben un prompt especial con reglas de clasificación distintas.")
    add_attr_row(tbl, "SYSTEM_PROMPT", "str",
                 "Prompt del sistema para Groq. Contiene reglas de clasificación (REAL, FALSO, SÁTIRA, ESTAFA, NO VERIFICABLE), banderas rojas a detectar, tipos de artículo y la instrucción de respuesta en JSON.")
    add_attr_row(tbl, "USER_PROMPT_BASE", "str",
                 "Template del prompt de usuario. Se llena con URL, dominio, contenido extraído y contexto de noticias similares. Define el formato exacto del JSON de respuesta.")
    add_attr_row(tbl, "SOCIAL_MEDIA_PROMPT", "str",
                 "Prompt especial para redes sociales con reglas de clasificación específicas por plataforma y ejemplos few-shot.")
    add_attr_row(tbl, "FEW_SHOT_EXAMPLES", "str",
                 "Ejemplos de clasificación (few-shot) para el LLM: REAL (noticia estándar), FALSO (afirmación médica falsa), NO VERIFICABLE (testimonio ambiguo).")
    add_attr_row(tbl, "USER_AGENTS", "List[str]",
                 "Lista de 5 User-Agent rotativos (Chrome macOS/Windows/Linux, Safari macOS, Firefox Windows). Se selecciona uno aleatorio por intento para evitar bloqueos por fingerprint.")
    add_attr_row(tbl, "BROWSER_HEADERS", "Dict[str, str]",
                 "Headers HTTP completos que imitan un navegador real (Accept, Accept-Language, Sec-Fetch-*, etc.).")
    add_attr_row(tbl, "LOGIN_PATTERNS", "List[str]",
                 "Patrones de texto en español e inglés que indican una página de inicio de sesión o bloqueo.")
    add_attr_row(tbl, "SCRAPING_PROXY", "str | None",
                 "Proxy opcional para scraping. Se lee de HTTP_PROXY o HTTPS_PROXY del entorno.")
    add_attr_row(tbl, "PW_USER_AGENTS", "List[str]",
                 "3 User-Agent específicos para Playwright (Chrome macOS, Chrome Windows, Firefox Windows).")
    doc.add_paragraph("")

    add_heading_styled(doc, "1.2.2 Funciones de Utilidad y Comunicación con Groq", 2)
    tbl = create_table(doc, ["Visibilidad", "Función", "Parámetros", "Retorno", "Descripción"])
    add_attr_row(tbl, "+", "get_groq_client()", "—", "Groq | None",
                 "Obtiene el cliente de Groq usando la API key de la variable de entorno GROQ_API_KEY. Retorna None si no hay API key configurada.")
    add_attr_row(tbl, "+", "call_groq(system_prompt, user_prompt)",
                 "system_prompt: str — instrucciones del sistema\nuser_prompt: str — contenido a analizar",
                 "str | None",
                 "Envía los prompts a Groq. Primero intenta con llama-3.3-70b-versatile; si falla, usa llama-3.1-8b-instant como fallback. Retorna la respuesta o None. temperature=0.1, response_format=json_object.")
    add_attr_row(tbl, "+", "get_domain(url)", "url: str — URL completa", "str",
                 "Extrae el dominio de una URL eliminando 'www.' Ej: 'https://www.jornada.com.mx/noticia' → 'jornada.com.mx'.")
    add_attr_row(tbl, "-", "_get_platform()", "—", "str",
                 "Detecta el SO (linux, windows, darwin) con sys.platform. Por defecto retorna 'linux'.")
    add_attr_row(tbl, "+", "parse_response(text)", "text: str — JSON raw de Groq", "dict | None",
                 "Parsea el JSON de Groq. Elimina bloques ``` si existen. Si falla el parseo, busca el primer '{' y último '}' como fallback.")
    doc.add_paragraph("")

    add_heading_styled(doc, "1.2.3 Estrategias de Scraping (Pipeline Multicapa)", 2)
    p = doc.add_paragraph(
        "Cada estrategia recibe una URL y retorna una tupla (respuesta, error). "
        "Si tiene éxito, el primer elemento es la respuesta HTTP. "
        "Si falla, retorna (None, mensaje_de_error)."
    )
    doc.add_paragraph("")

    add_heading_styled(doc, "_try_cloudscraper(url: str) → tuple", 3)
    p = doc.add_paragraph(
        "Primera estrategia. Usa cloudscraper para evadir Cloudflare imitando el handshake TLS de un navegador.\n"
        "Prueba 4 perfiles (Chrome Linux/Win, Firefox Linux/Win), cada uno con 2 intentos (8 total).\n"
        "Rota User-Agent aleatorio en cada intento. Usa proxy si SCRAPING_PROXY está configurado.\n"
        "Timeout: 30s por intento."
    )
    doc.add_paragraph("")

    add_heading_styled(doc, "_try_curl_cffi(url: str) → tuple", 3)
    p = doc.add_paragraph(
        "Segunda estrategia. Usa curl_cffi con impersonación de versiones específicas de navegador "
        "(chrome123, chrome120, safari17_0, chrome124). Cada versión imita el TLS fingerprint exacto.\n"
        "Rota User-Agent aleatorio. Usa proxy si está configurado.\n"
        "Timeout: 30s por intento."
    )
    doc.add_paragraph("")

    add_heading_styled(doc, "_try_requests(url: str) → tuple", 3)
    p = doc.add_paragraph(
        "Tercera estrategia. Usa requests con 2 intentos: primero verify=True, segundo verify=False.\n"
        "Rota User-Agent aleatorio. Usa proxy si está configurado.\n"
        "Timeout: 15s por intento. Es la más rápida pero la más fácil de bloquear."
    )
    doc.add_paragraph("")

    add_heading_styled(doc, "_try_playwright(url: str) → tuple", 3)
    p = doc.add_paragraph(
        "Cuarta estrategia. Lanza un navegador real (Firefox o Chromium) con perfil completo: "
        "viewport 1920x1080, locale es-MX, timezone America/Mexico_City.\n"
        "Prueba Firefox primero, luego Chromium. Cada motor se intenta 2 veces (4 intentos totales).\n"
        "Usa --no-sandbox para Linux/Render. Proxy configurable vía SCRAPING_PROXY.\n\n"
        "Flujo:\n"
        "1. Navega con wait_until='domcontentloaded'\n"
        "2. Espera hasta 35s (polling cada 1s) a que Cloudflare resuelva el challenge\n"
        "3. Verifica que el contenido no tenga indicadores de Cloudflare\n"
        "4. Verifica que el título sea contenido real\n\n"
        "Timeout por intento: ~95s (60s navegación + 35s polling)."
    )
    doc.add_paragraph("")

    add_heading_styled(doc, "_try_google_cache(url: str) → tuple", 3)
    p = doc.add_paragraph(
        "Quinta estrategia (fallback final). Obtiene el artículo desde Google Web Cache.\n"
        "Google ya descargó la página con su propia IP, por lo que evade completamente Cloudflare.\n\n"
        "URL: https://webcache.googleusercontent.com/search?q=cache:{url}\n"
        "2 intentos con requests estándar. Timeout: 20s por intento.\n\n"
        "NOTA: Solo funciona si Google tiene la página en caché (casi siempre cierto para sitios de noticias)."
    )
    doc.add_paragraph("")

    add_heading_styled(doc, "_http_get(url: str) → tuple (Orquestador)", 3)
    p = doc.add_paragraph(
        "Ejecuta las 5 estrategias en secuencia:\n\n"
        "cloudscraper → curl_cffi → requests → playwright → google_cache\n\n"
        "En cuanto una retorna éxito, la retorna inmediatamente. "
        "Si todas fallan, retorna None con los errores concatenados."
    )
    doc.add_paragraph("")

    add_heading_styled(doc, "1.2.4 Extracción de Contenido HTML", 2)
    add_heading_styled(doc, '_extract_from_html(html: str, domain: str = "") → dict', 3)
    p = doc.add_paragraph(
        "Parsea el HTML y extrae título, metadescripción y texto del artículo.\n\n"
        "Flujo:\n"
        "1. Detecta si el HTML viene de Google Cache y extrae el contenido interno\n"
        "2. Obtiene <title> y meta description\n"
        "3. Para Threads: extrae posts del JSON en <script type='application/json'>\n"
        "4. Elimina etiquetas no deseadas: script, style, nav, footer, aside, header, iframe, noscript\n"
        "5. Busca contenedor principal: <article> → <main> → <body>\n"
        "6. Para Instagram: usa meta description como texto principal\n"
        "7. Extrae párrafos <p> con >40 caracteres (hasta 50)\n"
        "8. Si texto < 500 caracteres, busca más párrafos (hasta 75)\n"
        "9. Si aún < 100 caracteres, usa get_text() filtrando líneas >60 caracteres\n\n"
        "Retorna: {content: str, title: str, article_text: str}"
    )
    doc.add_paragraph("")

    add_heading_styled(doc, "1.2.5 Manejo de Páginas Bloqueadas y Redes Sociales", 2)
    tbl = create_table(doc, ["Visibilidad", "Función", "Parámetros", "Retorno", "Descripción"])
    add_attr_row(tbl, "-", "_is_login_blocked_page(title, body)",
                 "title: str — título de la página\nbody: str — texto extraído",
                 "bool",
                 "Detecta si la página requiere inicio de sesión. Si body ≥ 800 caracteres, asume que no está bloqueada. Si no, busca patrones de login (mínimo 2 coincidencias).")
    add_attr_row(tbl, "-", "_extract_facebook_post_id(url)",
                 "url: str — URL de Facebook",
                 "str | None",
                 "Extrae el ID de un post de Facebook. Formatos: /{username}/posts/{id}, /videos/{id}, /photos/{id}, story.php?story_fbid={id}, photo.php?fbid={id}.")
    add_attr_row(tbl, "-", "_try_facebook_graph_api(url)",
                 "url: str — URL de Facebook",
                 "dict | None",
                 "Obtiene contenido de un post público de Facebook usando la Graph API. Requiere FACEBOOK_APP_ID y FACEBOOK_APP_SECRET en .env.")
    doc.add_paragraph("")

    add_heading_styled(doc, "1.2.6 Funciones Principales del Módulo", 2)

    add_heading_styled(doc, "scrape_url(url: str) → dict", 3)
    p = doc.add_paragraph(
        "Orquesta el scraping completo de una URL.\n\n"
        "Flujo:\n"
        "1. Llama a _http_get() que prueba las 5 estrategias\n"
        "2. Si todas fallan, retorna {'error': ...}\n"
        "3. Extrae contenido con _extract_from_html()\n"
        "4. Si texto < 500 caracteres, reintenta con Playwright\n"
        "5. Detecta login blocking; si es Facebook, intenta Graph API\n"
        "6. Si Facebook sin credenciales, retorna instrucciones"
    )
    doc.add_paragraph("")

    add_heading_styled(doc, "analyze_url(url: str) → dict", 3)
    p = doc.add_paragraph(
        "Función principal de análisis. Punto de entrada desde app.py.\n\n"
        "Flujo:\n"
        "1. Verifica GROQ_API_KEY\n"
        "2. Llama a scrape_url()\n"
        "3. Extrae dominio y determina si es fuente creíble\n"
        "4. Busca noticias similares vía Google News RSS\n"
        "5. Construye prompt con contenido, dominio y contexto\n"
        "6. Si es red social, usa SOCIAL_MEDIA_PROMPT\n"
        "7. Envía a Groq y parsea la respuesta\n"
        "8. Si dominio creíble y Groq retornó FALSO, reclasifica a NO VERIFICABLE\n"
        "9. Retorna análisis completo con metadatos\n\n"
        "Retorna: {analysis: dict, title: str, article_text: str, domain: str, is_credible_source: bool}\n"
        "En caso de error: {error: str, status: int}"
    )
    doc.add_paragraph("")

    add_heading_styled(doc, "1.3 news_finder.py — Módulo de Búsqueda de Noticias", 1)

    p = doc.add_paragraph()
    run = p.add_run("Archivo: ")
    run.bold = True
    p.add_run("server/news_finder.py").font.name = "Consolas"
    doc.add_paragraph("")

    tbl = create_table(doc, ["Visibilidad", "Función", "Parámetros", "Retorno", "Descripción"])
    add_attr_row(tbl, "+", "find_similar_news(query, max_results)",
                 "query: str — título del artículo (mín. 5 caracteres)\nmax_results: int = 5 — máximo de resultados",
                 "List[dict]",
                 "Busca noticias relacionadas usando Google News RSS (hl=es, gl=MX). "
                 "Retorna [{title, url, published, source}]. En caso de error, retorna lista vacía.")
    doc.add_paragraph("")


def add_frontend_section(doc):
    add_heading_styled(doc, "2. Frontend (React/TypeScript)", 0)
    doc.add_paragraph(
        "El frontend es una SPA construida con React 18, TypeScript y Vite. "
        "Arquitectura de componentes funcionales con estado local en UrlInput y "
        "estado global en App mediante hooks (useState, useCallback, useMemo)."
    )
    doc.add_paragraph("")

    add_heading_styled(doc, "2.1 Interfaces y Tipos Compartidos", 1)

    p = doc.add_paragraph()
    run = p.add_run("Nota: ")
    run.bold = True
    run.font.color.rgb = NOTE_COLOR
    p.add_run(
        "TypeScript no genera código en tiempo de ejecución. Las interfaces se "
        "definen en archivos .tsx y se usan para type-checking en compilación. "
        "No hay clases en el frontend — todos los componentes son funciones."
    )
    doc.add_paragraph("")

    add_heading_styled(doc, "2.1.1 Lang (type alias)", 2)
    p = doc.add_paragraph()
    p.add_run("Definida en: ").bold = True
    p.add_run("src/App.tsx:10").font.name = "Consolas"
    doc.add_paragraph("'es' | 'en' — Tipo unión que representa los idiomas soportados.")
    doc.add_paragraph("")

    add_heading_styled(doc, "2.1.2 Analysis (interface)", 2)
    p = doc.add_paragraph()
    p.add_run("Definida en: ").bold = True
    p.add_run("src/App.tsx:12").font.name = "Consolas"
    doc.add_paragraph("")
    tbl = create_table(doc, ["Propiedad", "Tipo", "Descripción"])
    add_attr_row(tbl, "verdict", "string", "Veredicto del análisis: REAL, FALSO, SÁTIRA, ESTAFA o NO VERIFICABLE")
    add_attr_row(tbl, "confidence_score", "number", "Puntuación de confianza del 0 al 100")
    add_attr_row(tbl, "summary", "string", "Resumen neutral del artículo en 2-3 oraciones")
    add_attr_row(tbl, "extracted_claims?", "string[]", "Afirmaciones principales extraídas (opcional)")
    add_attr_row(tbl, "reasoning", "string[]", "Razonamiento detallado detrás del veredicto")
    add_attr_row(tbl, "red_flags", "string[]", "Banderas rojas detectadas")
    add_attr_row(tbl, "positive_signals", "string[]", "Señales positivas detectadas")
    add_attr_row(tbl, "article_type?", "string",
                 "Tipo de artículo: informativa, comercial, opinion, clickbait, denuncia (opcional)")
    add_attr_row(tbl, "is_scam?", "boolean", "Indicador de posible estafa (opcional)")
    doc.add_paragraph("")

    add_heading_styled(doc, "2.1.3 NewsItem (interface)", 2)
    p = doc.add_paragraph()
    p.add_run("Definida en: ").bold = True
    p.add_run("src/App.tsx:24 y src/components/SimilarNews.tsx:3").font.name = "Consolas"
    doc.add_paragraph("")
    tbl = create_table(doc, ["Propiedad", "Tipo", "Descripción"])
    add_attr_row(tbl, "title", "string", "Título de la noticia")
    add_attr_row(tbl, "url", "string", "URL de la noticia")
    add_attr_row(tbl, "published", "string", "Fecha de publicación en formato raw")
    add_attr_row(tbl, "source", "string", "Nombre de la fuente o medio")
    doc.add_paragraph("")

    add_heading_styled(doc, "2.1.4 ApiResponse (interface)", 2)
    p = doc.add_paragraph()
    p.add_run("Definida en: ").bold = True
    p.add_run("src/App.tsx:31").font.name = "Consolas"
    doc.add_paragraph("")
    tbl = create_table(doc, ["Propiedad", "Tipo", "Descripción"])
    add_attr_row(tbl, "analysis", "Analysis | null", "Resultado del análisis o null si hay error")
    add_attr_row(tbl, "similar_news", "NewsItem[]", "Noticias similares encontradas")
    add_attr_row(tbl, "url_analyzed", "string", "URL que fue analizada")
    add_attr_row(tbl, "article_text", "string", "Texto extraído del artículo (hasta 2000 caracteres)")
    add_attr_row(tbl, "domain", "string", "Dominio extraído de la URL analizada")
    add_attr_row(tbl, "is_credible_source", "boolean", "Indica si el dominio está en la lista de medios reconocidos")
    add_attr_row(tbl, "error", "string | null", "Mensaje de error o null si la operación fue exitosa")
    doc.add_paragraph("")

    add_heading_styled(doc, "2.2 Componentes React", 1)

    add_heading_styled(doc, "2.2.1 App (App.tsx) — Componente Raíz", 2)
    p = doc.add_paragraph(
        "Componente principal que orquesta el estado global, la comunicación con el backend "
        "y el renderizado de todos los componentes hijos."
    )
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.add_run("Props: ").bold = True
    p.add_run("Ninguna (es el componente raíz, no recibe props).").font.name = "Consolas"
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.add_run("Estado local (useState):").bold = True
    doc.add_paragraph("")
    tbl = create_table(doc, ["Variable", "Tipo", "Valor Inicial", "Setter", "Descripción"])
    add_attr_row(tbl, "lang", "Lang ('es' | 'en')", "'es'", "setLang",
                 "Idioma actual de la interfaz. Controla TRANSLATIONS y se pasa a todos los hijos.")
    add_attr_row(tbl, "loading", "boolean", "false", "setLoading",
                 "Indica si hay una solicitud de análisis en curso.")
    add_attr_row(tbl, "result", "ApiResponse | null", "null", "setResult",
                 "Respuesta del backend tras un análisis exitoso.")
    add_attr_row(tbl, "error", "string | null", "null", "setError",
                 "Mensaje de error si el análisis falla.")
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.add_run("Valores derivados:").bold = True
    doc.add_paragraph("")
    tbl = create_table(doc, ["Variable", "Tipo", "Descripción"])
    add_attr_row(tbl, "tx", "Record<string, string>", "Traducciones para el idioma actual (TRANSLATIONS[lang])")
    add_attr_row(tbl, "analysis", "Analysis | null", "Resultado del análisis (result?.analysis ?? null)")
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.add_run("Callbacks (useCallback):").bold = True
    doc.add_paragraph("")
    tbl = create_table(doc, ["Nombre", "Firma", "Dependencias", "Descripción"])
    add_attr_row(tbl, "handleToggleLang", "() => void", "[]",
                 "Alterna el idioma entre 'es' y 'en'.")
    add_attr_row(tbl, "handleClear", "() => void", "[]",
                 "Limpia el resultado y el error actual.")
    add_attr_row(tbl, "handleAnalyze", "(url: string) => Promise<void>", "[lang, tx]",
                 "Ejecuta el análisis de una URL. AbortController con timeout de 60s, "
                 "POST a /analyze, maneja errores de red y timeout.")
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.add_run("Valores memorizados (useMemo):").bold = True
    doc.add_paragraph("")
    tbl = create_table(doc, ["Nombre", "Tipo", "Dependencias", "Descripción"])
    add_attr_row(tbl, "adjustedVerdict", "string | null", "[analysis, result]",
                 "Ajusta el veredicto: si score < 50 → FALSO, si score ≤ 69 → DUDOSO (excepto redes sociales).")
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.add_run("Componentes hijos que renderiza:").bold = True
    doc.add_paragraph("")
    tbl = create_table(doc, ["Componente", "Tipo de Carga", "Props"])
    add_attr_row(tbl, "LanguageToggle", "Eager (directa)", "lang, onToggle={handleToggleLang}")
    add_attr_row(tbl, "UrlInput", "Eager (directa)", "lang, loading, onAnalyze={handleAnalyze}, onClear={handleClear}")
    add_attr_row(tbl, "VerdictDisplay", "Eager (condicional)", "verdict, originalVerdict, lang")
    add_attr_row(tbl, "ConfidenceBar", "Eager (condicional)", "score, lang")
    add_attr_row(tbl, "RedFlags", "Eager (condicional)", "redFlags, positiveSignals, lang")
    add_attr_row(tbl, "SimilarNews", "Lazy (React.lazy + Suspense)", "news, lang")
    doc.add_paragraph("")

    add_heading_styled(doc, "2.2.2 UrlInput (UrlInput.tsx) — Input de URL", 2)
    p = doc.add_paragraph(
        "Formulario para ingresar la URL a analizar. Maneja su propio estado local "
        "para el valor del input y callbacks de submit/clear."
    )
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.add_run("Props (Props local):").bold = True
    doc.add_paragraph("")
    tbl = create_table(doc, ["Propiedad", "Tipo", "Descripción"])
    add_attr_row(tbl, "lang", "'es' | 'en'", "Idioma para textos del formulario")
    add_attr_row(tbl, "loading", "boolean", "Deshabilita el input y muestra estado de carga")
    add_attr_row(tbl, "onAnalyze", "(url: string) => void", "Callback al hacer submit con la URL")
    add_attr_row(tbl, "onClear?", "() => void", "Callback opcional al limpiar")
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.add_run("Estado local:").bold = True
    doc.add_paragraph("")
    tbl = create_table(doc, ["Variable", "Tipo", "Valor Inicial", "Descripción"])
    add_attr_row(tbl, "url", "string", "''", "Valor actual del input de URL")
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.add_run("Callbacks (useCallback):").bold = True
    doc.add_paragraph("")
    tbl = create_table(doc, ["Nombre", "Firma", "Dependencias", "Descripción"])
    add_attr_row(tbl, "handleChange", "(e: React.ChangeEvent<HTMLInputElement>) => void", "[]",
                 "Actualiza el estado url con el valor del input.")
    add_attr_row(tbl, "handleClear", "() => void", "[onClear]",
                 "Limpia el input y llama a onClear externo.")
    add_attr_row(tbl, "handleSubmit", "(e: React.FormEvent) => void", "[url, loading, onAnalyze]",
                 "Previene el envío del form, valida que url no esté vacía y no esté cargando, y llama a onAnalyze.")
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.add_run("Subcomponente interno: ").bold = True
    p.add_run("LoadingSegments").font.name = "Consolas"
    doc.add_paragraph(
        "Componente memoizado sin props. Renderiza 10 divs animados (segmentos de barra de carga) "
        "con animationDelay escalonado para un efecto de barrido."
    )
    doc.add_paragraph("")

    add_heading_styled(doc, "2.2.3 VerdictDisplay (VerdictDisplay.tsx) — Visualización de Veredicto", 2)
    p = doc.add_paragraph(
        "Muestra el veredicto del análisis con colores y sombras según el tipo de veredicto."
    )
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.add_run("Props (Props local):").bold = True
    doc.add_paragraph("")
    tbl = create_table(doc, ["Propiedad", "Tipo", "Descripción"])
    add_attr_row(tbl, "verdict", "string", "Texto del veredicto (REAL, FALSO, SÁTIRA, etc.)")
    add_attr_row(tbl, "originalVerdict?", "string", "Veredicto original de Groq antes de ajustes")
    add_attr_row(tbl, "lang", "'es' | 'en'", "Idioma para los subtítulos")
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.add_run("Valores memorizados (useMemo):").bold = True
    doc.add_paragraph("")
    tbl = create_table(doc, ["Nombre", "Tipo", "Dependencias", "Descripción"])
    add_attr_row(tbl, "containerStyle", "React.CSSProperties", "[cfg]",
                 "Estilo del contenedor: borde, background, clipPath y boxShadow según el veredicto.")
    add_attr_row(tbl, "wordStyle", "React.CSSProperties", "[cfg]",
                 "Estilo del texto: fuente Orbitron, tamaño 2.5rem, color y textShadow según el veredicto.")
    doc.add_paragraph("")

    add_heading_styled(doc, "2.2.4 ConfidenceBar (ConfidenceBar.tsx) — Barra de Confianza", 2)
    p = doc.add_paragraph(
        "Barra de 20 segmentos iluminados. Color según rango: rojo (0-40), naranja (41-69), "
        "cian (70-89), verde (90-100)."
    )
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.add_run("Props (Props local):").bold = True
    doc.add_paragraph("")
    tbl = create_table(doc, ["Propiedad", "Tipo", "Descripción"])
    add_attr_row(tbl, "score", "number", "Puntuación de confianza (0-100)")
    add_attr_row(tbl, "lang", "'es' | 'en'", "Idioma para la etiqueta")
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.add_run("Función de utilidad interna:").bold = True
    doc.add_paragraph("")
    tbl = create_table(doc, ["Nombre", "Firma", "Descripción"])
    add_attr_row(tbl, "getColorKey", "(score: number) => string",
                 "Retorna 'red' si score ≤ 40, 'orange' si ≤ 69, 'cyan' si ≤ 89, 'green' si ≥ 90.")
    doc.add_paragraph("")

    add_heading_styled(doc, "2.2.5 RedFlags (RedFlags.tsx) — Banderas Rojas y Señales Positivas", 2)
    p = doc.add_paragraph(
        "Dos paneles: alertas detectadas (rojo) y señales positivas (cian). "
        "Cada ítem se renderiza como un pill."
    )
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.add_run("Props (Props local):").bold = True
    doc.add_paragraph("")
    tbl = create_table(doc, ["Propiedad", "Tipo", "Descripción"])
    add_attr_row(tbl, "redFlags", "string[]", "Alertas y banderas rojas")
    add_attr_row(tbl, "positiveSignals", "string[]", "Señales positivas")
    add_attr_row(tbl, "lang", "'es' | 'en'", "Idioma para los títulos de los paneles")
    doc.add_paragraph("")

    add_heading_styled(doc, "2.2.6 SimilarNews (SimilarNews.tsx) — Noticias Similares", 2)
    p = doc.add_paragraph(
        "Cuadrícula de tarjetas con noticias relacionadas. "
        "Carga diferida (React.lazy). Cada tarjeta muestra fuente, título y fecha."
    )
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.add_run("Props (Props local):").bold = True
    doc.add_paragraph("")
    tbl = create_table(doc, ["Propiedad", "Tipo", "Descripción"])
    add_attr_row(tbl, "news", "NewsItem[]", "Lista de noticias similares")
    add_attr_row(tbl, "lang", "'es' | 'en'", "Idioma para el título de la sección")
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.add_run("Función de utilidad interna:").bold = True
    doc.add_paragraph("")
    tbl = create_table(doc, ["Nombre", "Firma", "Descripción"])
    add_attr_row(tbl, "formatDate", "(raw: string) => string",
                 "Convierte fecha raw a formato localizado es-MX (ej: '13 jul 2026'). "
                 "Si falla, retorna el string original.")
    doc.add_paragraph("")

    add_heading_styled(doc, "2.2.7 LanguageToggle (LanguageToggle.tsx) — Cambio de Idioma", 2)
    p = doc.add_paragraph(
        "Botón que alterna entre español e inglés. Muestra el código del idioma opuesto "
        "(si lang='es' muestra 'EN', viceversa)."
    )
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.add_run("Props (Props local):").bold = True
    doc.add_paragraph("")
    tbl = create_table(doc, ["Propiedad", "Tipo", "Descripción"])
    add_attr_row(tbl, "lang", "'es' | 'en'", "Idioma actual")
    add_attr_row(tbl, "onToggle", "() => void", "Callback al hacer clic")
    doc.add_paragraph("")


def add_dependency_section(doc):
    add_heading_styled(doc, "3. Diagrama de Dependencias entre Módulos", 0)
    doc.add_paragraph(
        "Arquitectura cliente-servidor con comunicación vía HTTP/JSON. "
        "A continuación se describe el flujo completo desde que el usuario ingresa una URL hasta que ve el resultado."
    )
    doc.add_paragraph("")

    add_heading_styled(doc, "3.1 Backend → Backend", 2)
    tbl = create_table(doc, ["Módulo Origen", "Módulo Destino", "Tipo de Conexión", "Descripción"])
    add_attr_row(tbl, "app.py", "analyzer.py", "importación directa", "Llama a analyze_url() en la ruta POST /analyze")
    add_attr_row(tbl, "app.py", "news_finder.py", "importación directa", "Llama a find_similar_news() después del análisis")
    add_attr_row(tbl, "analyzer.py", "news_finder.py", "importación dentro de función", "Obtiene contexto de fuentes similares en analyze_url()")
    add_attr_row(tbl, "analyzer.py", "Groq API", "HTTP externo", "call_groq() envía prompts a la API de Groq")
    add_attr_row(tbl, "analyzer.py", "Google News RSS", "HTTP externo", "news_finder.py consulta el feed RSS")
    add_attr_row(tbl, "analyzer.py", "Google Web Cache", "HTTP externo", "_try_google_cache() fallback anti-Cloudflare")
    add_attr_row(tbl, "analyzer.py", "URL objetivo", "HTTP externo", "Las 5 estrategias de scraping obtienen el HTML")
    doc.add_paragraph("")

    add_heading_styled(doc, "3.2 Frontend → Backend", 2)
    tbl = create_table(doc, ["Componente Origen", "Endpoint", "Método HTTP", "Descripción"])
    add_attr_row(tbl, "App.tsx", "/analyze", "POST",
                 "Envía {url: string} y recibe ApiResponse con el análisis completo.")
    add_attr_row(tbl, "Render (monitoreo)", "/health", "GET",
                 "Health check. Retorna {'status': 'ok'}.")
    add_attr_row(tbl, "Navegador", "/<path>", "GET",
                 "Sirve archivos estáticos del frontend. Las rutas SPA redirigen a index.html.")
    doc.add_paragraph("")

    add_heading_styled(doc, "3.3 Frontend → Frontend (Jerarquía de Componentes)", 2)
    tbl = create_table(doc, ["Componente Padre", "Componente Hijo", "Tipo de Relación"])
    add_attr_row(tbl, "main.tsx", "App", "Renderizado (ReactDOM.createRoot)")
    add_attr_row(tbl, "App", "LanguageToggle", "Composición directa")
    add_attr_row(tbl, "App", "UrlInput", "Composición directa (recibe callbacks)")
    add_attr_row(tbl, "App", "VerdictDisplay", "Composición condicional (solo si hay analysis)")
    add_attr_row(tbl, "App", "ConfidenceBar", "Composición condicional")
    add_attr_row(tbl, "App", "RedFlags", "Composición condicional")
    add_attr_row(tbl, "App", "SimilarNews", "Composición lazy (React.lazy + Suspense)")
    add_attr_row(tbl, "UrlInput", "LoadingSegments", "Composición condicional (solo si loading=true)")
    doc.add_paragraph("")

    add_heading_styled(doc, "3.4 Flujo de Datos Completo", 2)
    doc.add_paragraph(
        "1. Usuario ingresa URL en UrlInput\n"
        "2. UrlInput llama a onAnalyze(url) → handleAnalyze en App\n"
        "3. App.setLoading(true), hace POST /analyze\n"
        "4. Flask (app.py) recibe la solicitud y llama a analyze_url(url)\n"
        "5. analyze_url() → scrape_url(url) → _http_get(url) prueba 5 estrategias\n"
        "6. _extract_from_html() extrae título y texto del HTML\n"
        "7. analyze_url() busca noticias similares vía find_similar_news()\n"
        "8. Construye prompt y llama a call_groq()\n"
        "9. Groq retorna JSON con veredicto, confianza, razonamiento\n"
        "10. Reclasifica si es necesario (dominio creíble + FALSO → NO VERIFICABLE)\n"
        "11. Flask arma ApiResponse y la retorna como JSON\n"
        "12. App.setResult(data), loading=false\n"
        "13. React renderiza: VerdictDisplay, ConfidenceBar, RedFlags, SimilarNews"
    )


def add_usability_section(doc):
    add_heading_styled(doc, "4. Estudio de Usabilidad (Herramientas de acuerdo a las necesidades del levantamiento)", 0)
    doc.add_paragraph(
        "Este apartado describe las herramientas tecnológicas seleccionadas para cada módulo "
        "del sistema VERIFEX, justificando su elección con base en los requisitos funcionales "
        "y no funcionales identificados durante el levantamiento de requerimientos. "
        "Cada herramienta fue evaluada según criterios de idoneidad, rendimiento, "
        "curva de aprendizaje, comunidad y costo."
    )
    doc.add_paragraph("")

    add_heading_styled(doc, "4.1 Requerimientos del Levantamiento", 1)
    doc.add_paragraph(
        "El levantamiento de requerimientos identificó las siguientes necesidades clave que "
        "determinaron la selección de herramientas:"
    )
    doc.add_paragraph("")
    tbl = create_table(doc, ["Necesidad", "Requerimiento asociado", "Herramienta seleccionada", "Justificación"])
    add_attr_row(tbl, "Analizar credibilidad de noticias desde una URL",
                 "HU-01, HU-02, HU-18",
                 "Flask + Playwright + cloudscraper + curl_cffi + requests",
                 "Se requiere un backend ligero que procese URLs robustamente. "
                 "La pila de scraping multicapa garantiza extracción exitosa incluso con "
                 "protecciones Cloudflare, bloqueos por IP o JavaScript pesado.")
    add_attr_row(tbl, "Clasificar contenido con IA sin GPU local",
                 "HU-03, HU-19, HU-20",
                 "Groq API (llama-3.3-70b-versatile + fallback llama-3.1-8b-instant)",
                 "Groq ofrece inferencia ultrarrápida en la nube sin necesidad de hardware "
                 "especializado. El fallback garantiza disponibilidad. Temperature baja (0.1) y "
                 "response_format='json_object' aseguran respuestas consistentes y parseables.")
    add_attr_row(tbl, "Interfaz intuitiva y responsive",
                 "HU-11, HU-12, HU-15, HU-27, HU-28, HU-29",
                 "React 18 + TypeScript + Vite + Tailwind CSS",
                 "React permite componentes reutilizables con estado local. TypeScript añade "
                 "type safety. Vite optimiza builds y recarga en caliente. Tailwind agiliza el "
                 "diseño responsive con utilidades directas en JSX.")
    add_attr_row(tbl, "Búsqueda de contexto adicional",
                 "HU-08, HU-22",
                 "Google News RSS (feedparser)",
                 "API gratuita, sin autenticación, devuelve resultados en español de México "
                 "(hl=es, gl=MX). Integración directa vía feedparser sin dependencias pesadas.")
    add_attr_row(tbl, "Comunicación frontend-backend segura",
                 "HU-16, HU-17, HU-24, HU-25",
                 "Fetch API + AbortController + Flask-CORS",
                 "Fetch nativo del navegador evita librerías extras. AbortController maneja "
                 "timeouts de 60s. Flask-CORS permite origenes cruzados en desarrollo y producción.")
    add_attr_row(tbl, "Despliegue en la nube sin运维",
                 "HU-23, DEP-01, DEP-02, DEP-03",
                 "Render (Web Service + Static Site) + Gunicorn",
                 "Render ofrece despliegue gratuito con build automático, SSL y dominio público. "
                 "Gunicorn sirve Flask con múltiples workers para producción.")
    add_attr_row(tbl, "Pruebas automatizadas",
                 "HU-34, TEST-01, TEST-02, TEST-03",
                 "pytest + Vitest + Testing Library",
                 "pytest para backend (27 tests), Vitest + Testing Library para frontend "
                 "(52 tests en 7 archivos). Cobertura de integración completa.")
    doc.add_paragraph("")

    add_heading_styled(doc, "4.2 Criterios de Selección", 1)
    doc.add_paragraph(
        "Cada herramienta fue evaluada bajo los siguientes criterios antes de ser incorporada al proyecto:"
    )
    doc.add_paragraph("")

    tbl = create_table(doc, ["Criterio", "Descripción", "Peso"])
    add_attr_row(tbl, "Idoneidad funcional",
                 "La herramienta resuelve directamente la necesidad identificada en el levantamiento. "
                 "Se priorizaron herramientas con propósito específico sobre soluciones genéricas.",
                 "Alto")
    add_attr_row(tbl, "Rendimiento",
                 "Tiempo de respuesta, consumo de recursos y escalabilidad. "
                 "Se midió latencia de Groq API (~2-5s por análisis), tiempo de scraping "
                 "(~10-30s con fallbacks) y tamaño de build de producción (~200KB JS).",
                 "Alto")
    add_attr_row(tbl, "Madurez y comunidad",
                 "Versión estable, documentación actualizada, número de usuarios y "
                 "frecuencia de actualizaciones. Todas las herramientas seleccionadas "
                 "tienen más de 5 años de desarrollo activo.",
                 "Medio")
    add_attr_row(tbl, "Costo",
                 "Las herramientas seleccionadas son gratuitas (open source) o tienen "
                 "tier gratuitos generosos. Groq API ofrece crédito gratuito inicial; "
                 "Render tiene tier free con limitaciones razonables.",
                 "Medio")
    add_attr_row(tbl, "Curva de aprendizaje",
                 "Tiempo estimado para que un desarrollador con experiencia web básica "
                 "sea productivo. Flask y React son ampliamente conocidos; Playwright "
                 "requiere configuración adicional (instalación de navegadores).",
                 "Bajo")
    doc.add_paragraph("")

    add_heading_styled(doc, "4.3 Alternativas Evaluadas y Descartadas", 1)
    doc.add_paragraph(
        "Durante el levantamiento se evaluaron las siguientes alternativas que fueron descartadas "
        "por no cumplir con los criterios de selección:"
    )
    doc.add_paragraph("")
    tbl = create_table(doc, ["Herramienta", "Módulo", "Motivo de descarte"])
    add_attr_row(tbl, "Django + Django REST Framework",
                 "Backend",
                 "Sobredimensionado para un API con solo 2 endpoints. Flask es más ligero, "
                 "tiene menor curva de aprendizaje y arranca en segundos.")
    add_attr_row(tbl, "Next.js / Nuxt.js",
                 "Frontend",
                 "Renderizado SSR innecesario para una SPA de una sola vista. "
                 "React + Vite es más simple y el build estático es suficiente.")
    add_attr_row(tbl, "Selenium",
                 "Scraping",
                 "Más pesado que Playwright, requiere WebDriver y es más lento. "
                 "Playwright tiene mejor API para esperar condiciones y manejar Cloudflare.")
    add_attr_row(tbl, "OpenAI API / Claude API",
                 "Clasificación IA",
                 "Mayor latencia y costo por consulta. Groq ofrece mejor rendimiento "
                 "para tareas de clasificación estructurada con respuesta JSON.")
    add_attr_row(tbl, "Heroku / Railway",
                 "Despliegue",
                 "Heroku ya no tiene tier gratuito. Railway tiene limitaciones de memoria "
                 "que causaron fallos en el despliegue del scraper con Playwright.")
    add_attr_row(tbl, "Axios",
                 "HTTP frontend",
                 "Fetch API nativo es suficiente para una sola llamada POST. "
                 "Axios añadiría ~14KB innecesarios al bundle.")
    doc.add_paragraph("")

    add_heading_styled(doc, "4.4 Mapeo de Herramientas a Requerimientos Funcionales", 1)
    doc.add_paragraph(
        "La siguiente tabla resume la correspondencia entre los requerimientos funcionales "
        "(HU) y las herramientas que los implementan:"
    )
    doc.add_paragraph("")
    tbl = create_table(doc, ["HU", "Descripción", "Herramienta(s)", "Módulo"])
    add_attr_row(tbl, "HU-01, HU-02", "Pegar URL y ver veredicto", "React + Flask + Groq", "Frontend/Backend")
    add_attr_row(tbl, "HU-03", "Nivel de confianza (20 niveles)", "React (ConfidenceBar)", "Frontend")
    add_attr_row(tbl, "HU-04", "Banderas rojas y señales positivas", "React (RedFlags) + Groq", "Frontend/Backend")
    add_attr_row(tbl, "HU-05, HU-06", "Resumen y afirmaciones con IA", "Groq API + analyzer.py", "Backend")
    add_attr_row(tbl, "HU-07", "Análisis por categoría", "Groq API (prompt engineering)", "Backend")
    add_attr_row(tbl, "HU-08, HU-22", "Noticias similares", "Google News RSS + news_finder.py", "Backend")
    add_attr_row(tbl, "HU-09", "Tipo de artículo", "Groq API (clasificación)", "Backend")
    add_attr_row(tbl, "HU-10", "Alerta de estafa", "Groq API (is_scam) + React", "Frontend/Backend")
    add_attr_row(tbl, "HU-11", "Cambio de idioma", "React (LanguageToggle)", "Frontend")
    add_attr_row(tbl, "HU-12, HU-26, HU-27, HU-28, HU-29", "Interfaz cyberpunk responsive", "React + Tailwind + CSS custom", "Frontend")
    add_attr_row(tbl, "HU-16", "Sin almacenamiento de datos", "Flask (sin BD) + Fetch API", "Arquitectura")
    add_attr_row(tbl, "HU-17", "Conexión segura HTTPS", "Render (SSL automático)", "Despliegue")
    add_attr_row(tbl, "HU-18", "Scraper robusto", "cloudscraper + curl_cffi + requests + Playwright + Google Cache", "Backend")
    add_attr_row(tbl, "HU-19, HU-20", "Clasificación con IA y prompt engineering", "Groq API + analyzer.py", "Backend")
    add_attr_row(tbl, "HU-21", "Override de dominios confiables", "CREDIBLE_DOMAINS + analyzer.py", "Backend")
    add_attr_row(tbl, "HU-23", "Despliegue zero downtime", "Render + Gunicorn", "Despliegue")
    add_attr_row(tbl, "HU-24", "Backend con /analyze y /health", "Flask + app.py", "Backend")
    add_attr_row(tbl, "HU-25", "Frontend con fetch + timeout", "React + Fetch API + AbortController", "Frontend")
    add_attr_row(tbl, "HU-30 a HU-34", "Documentación, diagramas, pruebas", "pytest + Vitest + UML + Markdown", "Documentación")
    doc.add_paragraph("")

    add_heading_styled(doc, "4.5 Conclusión del Estudio de Usabilidad", 1)
    doc.add_paragraph(
        "Las herramientas seleccionadas cumplen con los requerimientos del levantamiento "
        "al priorizar simplicidad, rendimiento y costo cero. La arquitectura resultante "
        "es mantenible, extensible y desplegable sin infraestructura compleja. "
        "Las alternativas descartadas fueron evaluadas objetivamente y no se descarta "
        "su incorporación futura si los requerimientos del proyecto evolucionan."
    )


def main():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    add_title_page(doc)
    add_backend_section(doc)
    doc.add_page_break()
    add_frontend_section(doc)
    doc.add_page_break()
    add_dependency_section(doc)
    doc.add_page_break()
    add_usability_section(doc)

    filepath = os.path.join(OUT, "Modulos_Programacion_VERIFEX.docx")
    doc.save(filepath)
    print(f"Documento generado: {filepath}")


if __name__ == "__main__":
    main()
