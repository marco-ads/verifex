#!/usr/bin/env python3
"""
Generador del Informe Técnico VERIFEX — Versión didáctica accesible
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Estilos base ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 5):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

def add_tip_box(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"  {text}")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x39, 0x49, 0xAB)
    return p

def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '1A237E')
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
            if ri % 2 == 0:
                set_cell_shading(cell, 'F5F5F5')
    return table

# ═══════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("INFORME TECNICO FORENSE")
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("VERIFEX: Analizador de Credibilidad de Noticias")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x39, 0x49, 0xAB)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Proyecto de Tesis — Deteccion de Noticias Falsas con Inteligencia Artificial")
run.font.size = Pt(12)
run.font.italic = True
run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Integrantes del Equipo:")
run.font.size = Pt(11)
run.font.bold = True

for name, role in [
    ("Marco Antonio Delgado Serrano", "Backend, IA, Scraping, Deploy, Pruebas"),
    ("Luis", "Frontend, Diseno UI/UX Cyberpunk, Responsive"),
    ("Ulises", "Documentacion de Tesis (Capitulos 1-5)"),
    ("Tony", "Diagramas UML, Documentacion Tecnica, Pruebas"),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{name} — {role}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x42, 0x42, 0x42)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Fecha: Julio 2026")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# TABLA DE CONTENIDOS (manual)
# ═══════════════════════════════════════════════════════════

doc.add_heading("Tabla de Contenidos", level=1)

toc_items = [
    "1. Arquitectura General y Metodologia de Trabajo",
    "   1.1 Que metodologia usamos y por que",
    "   1.2 En que sprint estamos",
    "   1.3 Estructura del equipo",
    "2. Stack Tecnologico y Librerias",
    "   2.1 Backend (la parte del servidor)",
    "   2.2 Frontend (la parte que ve el usuario)",
    "3. Motor de IA y Procesamiento de Lenguaje Natural",
    "   3.1 Como funciona la inteligencia artificial",
    "   3.2 El pipeline de inferencia (paso a paso)",
    "   3.3 Categorias de veredicto",
    "   3.4 Explicabilidad (por que la IA decide algo)",
    "4. Extraccion y Procesamiento del Texto",
    "   4.1 Flujo completo de datos",
    "   4.2 Scraping (como leemos paginas web)",
    "   4.3 Limpieza del texto",
    "   4.4 Persistencia (donde guardamos los datos)",
    "5. Comunicacion Frontend-Backend y Despliegue",
    "   5.1 Como viaja la informacion",
    "   5.2 Formato de la respuesta",
    "6. Base de Datos y Persistencia",
    "7. Technical Debt y Mejoras",
    "   7.1 Cuellos de botella",
    "   7.2 Problemas de seguridad",
    "   7.3 Recomendaciones concretas",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    for r in p.runs:
        r.font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 1. ARQUITECTURA GENERAL
# ═══════════════════════════════════════════════════════════

doc.add_heading("1. Arquitectura General y Metodologia de Trabajo", level=1)

doc.add_heading("1.1 Que metodologia usamos y por que", level=2)

doc.add_paragraph(
    "Metodologia identificada: Scrum adaptado con elementos de Kanban."
)

doc.add_paragraph(
    "Para entender esto, piensa en como se organiza un equipo de construction. "
    "Existen diferentes formas de trabajar:"
)

doc.add_paragraph(
    'Scrum: Es como si el equipo dividiera el proyecto en "sprints" (carreras cortas de 1-3 semanas). '
    "En cada sprint, el equipo se compromete a terminar un conjunto especifico de tareas. "
    "Al final del sprint, se revisa lo que se hizo y se planifica el siguiente sprint. "
    "Es como jugar futbol por tiempos: cada tiempo tiene un objetivo claro.",
    style='List Bullet'
)

doc.add_paragraph(
    'Kanban: Es como una cinta transportadora en una fabrica. '
    "Las tareas se mueven de izquierda a derecha a traves de columnas (Backlog, Por Hacer, En Proceso, etc.). "
    "No hay tiempos fijos, las tareas avanzan cuando hay espacio. "
    "Es como el trafico de un semaforo: las tareas avanzan cuando hay luz verde.",
    style='List Bullet'
)

doc.add_paragraph(
    "En VERIFEX usamos ambos: Scrum para organizar el proyecto en ciclos grandes, "
    "y Kanban para ver el estado actual de cada tarea.",
    style='List Bullet'
)

doc.add_paragraph(
    "Evidencia de Scrum en el codigo:"
)

doc.add_paragraph(
    'El archivo "Gantt_VERIFEX.xlsx" define 6 ciclos (Ciclo 0 a Ciclo 5), '
    "cada uno con multiples sprints. Por ejemplo, el Ciclo 3 (Desarrollo del Sistema) "
    "tiene 6 rondas de trabajo paralelo, cada una con tareas de 7 a 13 dias de duracion.",
    style='List Bullet'
)

doc.add_paragraph(
    'El archivo "Kanban_VERIFEX.xlsx" complementa con columnas WIP-limited '
    "(WIP = Work In Progress, o sea, trabajo en progreso). "
    "Por ejemplo, la columna 'Por Hacer' tiene limite de 5 tareas, "
    "y 'En Progreso' tiene limite de 3. Esto evita que el equipo tenga demasiadas tareas a la vez.",
    style='List Bullet'
)

doc.add_paragraph(
    "Hitos de version: El proyecto tiene 12 entregas intermedias (v0.1.0 hasta v1.2.0). "
    "Cada version representa una mejora concreta del sistema. "
    "Es como los niveles de un videojuego: cada nivel desbloquea nuevas funcionalidades.",
    style='List Bullet'
)

add_tip_box(doc,
    "Dato curioso: Los 'sprints' en Scrum se llaman asi por las carreras cortas de atletismo. "
    "La idea es correr rapido y enfocado en una direccion, descansar, y volver a correr."
)

doc.add_heading("1.2 En que sprint estamos", level=2)

doc.add_paragraph(
    "Segun el diagrama de Gantt, actualmente estamos en el Ciclo 5 "
    "(Resultados y Documentacion Final), que es el ultimo ciclo del proyecto."
)

doc.add_paragraph(
    "Las tareas activas en este momento son:"
)

for task in [
    "Analisis con URLs reales de medios mexicanos (Milenio, Reforma, Aristegui, etc.)",
    "Diseno responsive (que se vea bien en celular, tablet y escritorio)",
    "Redaccion del Capitulo 5 de la tesis (Resultados y Conclusiones)",
    "Pruebas de regresion (verificar que nada se rompio despues de los cambios recientes)",
    "Correccion de errores en timeouts y edge cases (casos especiales)",
]:
    doc.add_paragraph(task, style='List Bullet')

doc.add_paragraph(
    "Si contamos los sprints numericamente, estariamos en el Sprint 26 de 30 posibles. "
    "Esto significa que ya completamos aproximadamente el 87% del proyecto."
)

doc.add_heading("1.3 Estructura del equipo", level=2)

make_table(doc,
    ["Rol", "Nombre", "Responsabilidades", "Color en Kanban"],
    [
        ["Backend & IA", "Marco", "Servidor, inteligencia artificial, scraping, despliegue, pruebas", "Morado (#F3E5F5)"],
        ["Frontend", "Luis", "Interfaz de usuario, diseno cyberpunk, responsive", "Azul (#E3F2FD)"],
        ["Documentacion", "Ulises", "Redaccion de la tesis (Capitulos 1-5), Gantt, Kanban", "Verde (#E8F5E9)"],
        ["Pruebas & UML", "Tony", "Diagramas de diseño, documentacion tecnica, revision de logica", "Rojo (#FFEBEE)"],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 2. STACK TECNOLOGICO
# ═══════════════════════════════════════════════════════════

doc.add_heading("2. Stack Tecnologico y Librerias", level=1)

doc.add_paragraph(
    "El 'stack tecnologico' es como la caja de herramientas que usamos para construir el proyecto. "
    "Asi como un carpintero necesita martillo, sierra y clavos, nosotros necesitamos programas "
    "y librerias (conjuntos de codigo ya hecho por otros programadores) para construir VERIFEX."
)

doc.add_heading("2.1 Backend (la parte del servidor)", level=2)

doc.add_paragraph(
    "El backend es como la cocina de un restaurante: el usuario nunca la ve, "
    "pero ahi se prepara todo (el analisis de la noticia). "
    "El usuario solo recibe el plato servido (el resultado en pantalla)."
)

doc.add_paragraph(
    "El backend esta escrito en Python, que es un lenguaje de programacion muy popular "
    "porque es facil de leer y tiene muchas librerias disponibles."
)

make_table(doc,
    ["Libreria", "Que es (explicacion simple)", "Que hace en VERIFEX", "Como se conecta con el resto"],
    [
        ["Flask",
         "Un 'framework' (marco de trabajo) web minimalista. "
         "Es como los cimientos de una casa: te da la estructura basica para crear "
         "un servidor web sin complicarte con cosas tecnicas complejas.",
         "Crea el servidor que recibe las peticiones del frontend y devuelve respuestas. "
         "Tiene dos rutas principales: /analyze (para analizar URLs) y /health (para verificar que funciona).",
         "Se conecta al frontend via HTTP (el protocolo de internet). "
         "El frontend le manda una URL y el backend le devuelve el analisis en JSON."],

        ["flask-cors",
         "CORS significa 'Cross-Origin Resource Sharing' (Comparticion de Recursos entre Origenes). "
         "Es como un portero que decide que visitantes pueden entrar al edificio.",
         "Permite que el frontend (que vive en un dominio diferente) haga peticiones al backend "
         "sin que el navegador las bloquee.",
         "Sin este, el navegador pensaria que el frontend esta intentando acceder "
         "a datos de otro sitio web y bloquearia la peticion."],

        ["groq",
         "Es la libreria para conectarse a Groq API, que es un servicio en la nube "
         "que ejecuta modelos de inteligencia artificial (como un大脑 en la nube).",
         "Envia el texto de la noticia a Groq para que la IA lo analice y devuelva "
         "un veredicto (REAL, FALSO, etc.).",
         "Es el corazon del proyecto: sin Groq, no hay analisis. "
         "El backend le manda el texto y Groq responde con el veredicto."],

        ["cloudscraper",
         "Es una libreria que engana a los sitios web que usan Cloudflare "
         "(un sistema de seguridad que bloquea bots). "
         "Es como si te vistieras de delivery para entrar a un edificio.",
         "Intenta acceder a la pagina web de la noticia como si fuera un humano normal. "
         "Es la primera estrategia de scraping.",
         "Si funciona, el backend ya tiene el texto de la noticia y continua con el analisis. "
         "Si no funciona, intenta con la siguiente estrategia."],

        ["curl_cffi",
         "Es otra forma de hacer peticiones web, pero mas parecida a como lo hace un navegador real. "
         "Es como si usaras un disfraz de Chrome o Firefox para acceder a la pagina.",
         "Es la segunda estrategia de scraping. "
         "Impersona (hacese pasar por) diferentes navegadores (Chrome, Safari) "
         "para que la pagina piense que es un humano.",
         "Se usa cuando cloudscraper no funciona (por ejemplo, cuando la pagina tiene proteccion extra)."],

        ["requests",
         "La libreria mas basica para hacer peticiones HTTP en Python. "
         "Es como llamar por telefono: haces una pregunta y esperas la respuesta.",
         "Tercera estrategia de scraping. "
         "Es la forma mas simple de acceder a una pagina web.",
         "Es el fallback mas basico: si las estrategias anteriores fallan, "
         "intenta algo simple antes de usar Playwright."],

        ["Playwright",
         "Es como tener un navegador real controlado por una computadora. "
         "Abrir Firefox o Chromium (las versiones de codigo abierto de Chrome), "
         "navegar a la pagina, esperar a que cargue, y extraer el texto.",
         "Es la cuarta y ultima estrategia de scraping. "
         "Es la mas lenta pero la mas potente: puede manejar JavaScript, "
         "esperar a que la pagina cargue completamente, y manejar protecciones Cloudflare.",
         "Se usa como ultimo recurso. Puede tardar hasta 35 segundos "
         "pero es el que tiene mas posibilidades de funcionar."],

        ["BeautifulSoup + lxml",
         "BeautifulSoup es un 'parser' (analizador) de HTML. "
         "Es como un detective que mira el codigo fuente de una pagina web "
         "y extrae solo el texto importante, ignorando los anuncios, menus, etc.",
         "Despues de que cualquiera de las 4 estrategias de scraping trae el codigo HTML, "
         "BeautifulSoup lo procesa y extrae el titulo, descripcion y cuerpo del articulo.",
         "Se conecta con el backend: el texto extraido se envia a Groq para su analisis."],

        ["feedparser",
         "Es una libreria para leer 'feeds RSS' (formato que usan los sitios de noticias "
         "para publicar sus ultimas noticias automaticamente).",
         "En 'news_finder.py', busca noticias similares en Google News RSS "
         "usando el titulo del articulo como busqueda.",
         "Las noticias similares se agregan al contexto que se envia a Groq, "
         "para que la IA pueda comparar informacion de multiples fuentes."],

        ["python-dotenv",
         "Es como un cajon seguro donde guardas tus contrasenas y llaves. "
         "Lee un archivo llamado '.env' que contiene informacion sensible "
         "(como la clave de API de Groq) sin exponerla en el codigo.",
         "Carga la variable GROQ_API_KEY (la clave de acceso a la inteligencia artificial) "
         "desde el archivo server/.env.",
         "Es la forma segura de pasarle la clave de Groq al backend "
         "sin que quede escrita en el codigo fuente."],

        ["pytest",
         "Es un 'framework' de pruebas para Python. "
         "Es como un examen automatico: le haces preguntas al codigo "
         "y el te dice si esta bien o mal.",
         "Ejecuta 27 pruebas automaticas en server/test_analyzer.py "
         "para verificar que las funciones principales funcionan correctamente.",
         "Se ejecuta antes de cada despliegue para asegurar que no se rompio nada."],
    ]
)

add_tip_box(doc,
    "Analogia simple: Imagina que VERIFEX es un restaurante. "
    "Flask es el mostrador donde pides tu comida (la URL). "
    "cloudscraper/curl_cffi/requests/Playwright son los repartidores que van a buscar los ingredientes (el texto de la noticia). "
    "BeautifulSoup es el chef que limpia y prepara los ingredientes. "
    "Groq es el chef principal que cocina el plato (el analisis). "
    "python-dotenv es la caja fuerte donde guardas la llave del restaurante."
)

doc.add_heading("2.2 Frontend (la parte que ve el usuario)", level=2)

doc.add_paragraph(
    "El frontend es como la sala de un restaurante: es donde el usuario se sienta, "
    "ve el menu (la interfaz), pide su plato (introduce la URL) "
    "y recibe la comida servida (el resultado del analisis)."
)

doc.add_paragraph(
    "El frontend esta escrito en TypeScript (una version mejorada de JavaScript "
    "que detecta errores antes de que pasen) y usa React como libreria principal."
)

make_table(doc,
    ["Libreria", "Que es (explicacion simple)", "Que hace en VERIFEX", "Como se conecta"],
    [
        ["React 18",
         "Es la libreria mas popular para crear interfaces web. "
         "Es como un set de piezas de LEGO: tienes componentes (botones, tarjetas, graficos) "
         "que puedes combinar para crear la interfaz.",
         "Crea todos los componentes visuales: UrlInput (donde escribes la URL), "
         "VerdictDisplay (muestra el veredicto), ConfidenceBar (muestra la barra de confianza), "
         "RedFlags (alertas), SimilarNews (noticias similares), LanguageToggle (cambiar idioma).",
         "Se conecta al backend via 'fetch' (una funcion de JavaScript que hace peticiones HTTP). "
         "Cuando el usuario introduce una URL, React le manda un POST a /analyze."],

        ["TypeScript",
         "Es JavaScript pero con 'tipos' (etiquetas que le dicen al codigo "
         "que tipo de dato es cada variable: texto, numero, lista, etc.). "
         "Es como si le pusieras etiquetas a tus cajones para saber donde guarda la ropa.",
         "Define las estructuras de datos: la interfaz 'Analysis' describe exactamente "
         "que campos devuelve Groq (verdict, confidence_score, summary, etc.). "
         "Esto evita errores como intentar multiplicar un texto por un numero.",
         "TypeScript revisa el codigo antes de que se ejecute, "
         "como un editor que corrige ortografia antes de publicar un libro."],

        ["Vite",
         "Es un 'bundler' (empaquetador) y servidor de desarrollo. "
         "Es como una maquina empacadora: toma todos los archivos sueltos "
         "(HTML, CSS, JavaScript, imagenes) y los junta en un solo archivo optimizado.",
         "Cuando ejecutas 'npm run build', Vite toma todos los archivos de src/ "
         "y genera la carpeta dist/ con los archivos listos para produccion.",
         "El frontend se sirve desde dist/, que es donde Flask busca los archivos "
         "cuando el usuario accede a la pagina principal."],

        ["Tailwind CSS",
         "Es un 'framework' de CSS (el lenguaje que le da estilo a las paginas web). "
         "En vez de escribir codigo CSS largo, usas clases cortas como 'flex', 'p-4', 'bg-blue-500'. "
         "Es como usar prefabricados en vez de construir desde cero.",
         "Le da el estilo cyberpunk a toda la interfaz: colores neón (cyan, rojo, verde), "
         "fuentes monoespaciadas, bordes con clip-path, efectos de brillo (glow).",
         "Se integra con React: cada componente usa clases de Tailwind "
         "para verse como una terminal de ciencia ficcion."],

        ["Vitest + testing-library",
         "Vitest es un 'framework' de pruebas para JavaScript (similar a pytest pero para frontend). "
         "testing-library es una libreria que te ayuda a probar componentes React "
         "como si fueras un usuario real (haciendo clic, escribiendo, esperando resultados).",
         "Ejecuta 52 pruebas automaticas que verifican que cada componente "
         "funciona correctamente (UrlInput acepta URLs, ConfidenceBar muestra el puntaje, etc.).",
         "Se ejecuta en paralelo con las pruebas del backend para asegurar "
         "que toda la aplicacion funciona de punta a punta."],

        ["AbortController",
         "Es una funcion nativa de JavaScript (no es una libreria externa). "
         "Es como un temporizador de bomba: si la peticion tarda mas de 60 segundos, "
         "la cancela automaticamente.",
         "En el frontend, cuando el usuario analiza una URL, se crea un AbortController "
         "con timeout de 60 segundos. Si Groq tarda demasiado, se cancela y se muestra un error.",
         "Se conecta con el backend: si el backend no responde en 60 segundos, "
         "el frontend muestra 'El analisis tardo demasiado. Intenta de nuevo.'"],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 3. MOTOR DE IA
# ═══════════════════════════════════════════════════════════

doc.add_heading("3. Motor de IA y Procesamiento del Lenguaje Natural", level=1)

doc.add_paragraph(
    "Esta es la parte mas importante del proyecto: la Inteligencia Artificial "
    "que analiza las noticias y decide si son verdaderas o falsas."
)

doc.add_heading("3.1 Como funciona la inteligencia artificial", level=2)

doc.add_paragraph(
    "VERIFEX NO ejecuta ningun modelo de IA localmente. "
    "En vez de eso, usa un servicio externo llamado Groq API "
    "(como un experto al que le preguntas por telefono)."
)

doc.add_paragraph(
    "Groq API es como tener un grupo de expertos en la nube (servidores remotos) "
    "que pueden leer y entender texto. Nosotros le enviamos el texto de la noticia "
    "y ellos nos devuelven su opinion: si es verdadera, falsa, satira, etc."
)

doc.add_paragraph(
    "El modelo que usamos se llama llama-3.3-70b-versatile "
    "(si el nombre suena raro, es porque los modelos de IA suelen tener nombres curiosos). "
    "Es un modelo de lenguaje grande (LLM) con 70 mil millones de parametros. "
    "Para que te hagas una idea, es como un cerebro digital con 70 mil millones de conexiones."
)

doc.add_paragraph(
    "Si Groq falla, tenemos un fallback (plan B) con un modelo mas pequeno: "
    "llama-3.1-8b-instant. Es como tener un experto de respaldo por si el principal "
    "no esta disponible."
)

add_tip_box(doc,
    "Analogia: Imagina que le envias una carta a un experto en noticias. "
    "El experto la lee, analiza todas las palabras, compara con su conocimiento, "
    "y te responde: 'Esta noticia es verdadera porque...' o 'Esta noticia es falsa porque...'. "
    "Groq API es ese experto, pero digital y muy rapido."
)

doc.add_heading("3.2 El pipeline de inferencia (paso a paso)", level=2)

doc.add_paragraph(
    "El 'pipeline de inferencia' es como una linea de ensamblaje en una fabrica. "
    "Cada estacion hace una tarea diferente, y el producto final (el analisis) "
    "va pasando por cada estacion hasta estar listo."
)

doc.add_paragraph("Paso 1: El usuario introduce una URL")
doc.add_paragraph(
    "El usuario escribe la direccion de una noticia en el cuadro de texto "
    "y hace clic en el boton 'Analizar'. "
    "Esto envia la URL al servidor via una peticion HTTP POST a /analyze."
)

doc.add_paragraph("Paso 2: Scraping (leer la pagina web)")
doc.add_paragraph(
    "El servidor intenta acceder a la pagina web usando 4 estrategias diferentes "
    "(cloudscraper, curl_cffi, requests, Playwright). "
    "Es como si intentaras abrir una puerta con 4 llaves diferentes: "
    "si la primera no funciona, intentas con la segunda, y asi sucesivamente."
)

doc.add_paragraph("Paso 3: Extraccion del texto")
doc.add_paragraph(
    "Una vez que se accede a la pagina, BeautifulSoup analiza el codigo HTML "
    "y extrae solo el texto importante: titulo, descripcion y cuerpo del articulo. "
    "Es como un detective que busca las pistas importantes en una escena del crimen "
    "e ignora todo lo demas."
)

doc.add_paragraph("Paso 4: Construccion del prompt")
doc.add_paragraph(
    "El 'prompt' es el mensaje que le enviamos a la IA para que analice el texto. "
    "Es como armar una carta muy bien redactada: incluye instrucciones especificas, "
    "ejemplos de como queremos la respuesta, y el texto de la noticia."
)

doc.add_paragraph("Paso 5: Llamada a Groq API")
doc.add_paragraph(
    "El servidor envia el prompt a Groq via la libreria groq. "
    "Groq ejecuta el modelo llama-3.3-70b y devuelve una respuesta en formato JSON "
    "(un texto estructurado como una tabla)."
)

doc.add_paragraph("Paso 6: Post-proceso")
doc.add_paragraph(
    "El servidor verifica que la respuesta de Groq sea valida "
    "(que tenga todos los campos necesarios). Si algo falta, "
    "llama a Groq de nuevo para traducir el texto al ingles "
    "(usando la funcion translate_analysis)."
)

doc.add_paragraph("Paso 7: Respuesta al frontend")
doc.add_paragraph(
    "El servidor envia el analisis completo al frontend, "
    "que lo muestra en pantalla con colores, graficos y explicaciones."
)

add_tip_box(doc,
    "El pipeline completo dura entre 5 y 15 segundos, dependiendo de la velocidad "
    "de la pagina web y de Groq. El scraping es la parte mas lenta "
    "(puede tardar hasta 35 segundos si usa Playwright)."
)

doc.add_heading("3.3 Categorias de veredicto", level=2)

doc.add_paragraph(
    "La IA puede devolver 5 tipos de veredicto. Piensa en ellos como 5 etiquetas "
    "que le ponemos a cada noticia:"
)

make_table(doc,
    ["Veredicto", "Que significa (explicacion simple)", "Ejemplo"],
    [
        ["REAL",
         "La noticia es verdadera. Viene de un medio de comunicacion reconocido "
         "(como Milenio, Reforma, CNN) y reporta hechos de forma periodistica estandar. "
         "Es como un articulo normal que lees en el periodico.",
         "Un articulo de El Universal sobre una decision del Banco de Mexico."],

        ["FALSO",
         "La noticia es falsa. Contiene informacion que se puede DEMOSTRAR que es incorrecta. "
         "No es solo 'no puedo verificar', sino que hay pruebas de que es mentira. "
         "Es como cuando alguien dice que la Tierra es plana.",
         "Un articulo que dice 'Cientificos confirman que el limon cura el cancer' sin ningun estudio."],

        ["SÁTIRA",
         "La noticia es humor o parodia. No es informativa, es para reir. "
         "Es como los articulos de El Deforma o The Onion. "
         "Importante: no la marcamos como FALSO porque no esta intentando enganar.",
         "Un titular que dice 'Gobierno declara dia nacional de la siesta'."],

        ["ESTAFA",
         "La noticia es un fraude o phishing. Intenta robar tu dinero o datos personales. "
         "Es como esas cadenas de WhatsApp que te dicen 'Gana dinero rapido'.",
         "Una pagina que vende 'productos milagrosos' con resultados garantizados."],

        ["NO VERIFICABLE",
         "No hay suficiente informacion para decidir. La noticia hace afirmaciones serias "
         "pero sin fuentes verificables. Es como cuando alguien dice 'dicen que...' "
         "pero no dice quien lo dice.",
         "Un articulo que dice 'Fuentes anonimas revelan irregularidades' sin nombres ni pruebas."],
    ]
)

doc.add_heading("3.4 Explicabilidad (por que la IA decide algo)", level=2)

doc.add_paragraph(
    "VERIFEX NO usa tecnicas como LIME o SHAP "
    "(que son metodos para explicar por que un modelo de IA tomo una decision). "
    "En vez de eso, la explicabilidad viene del propio modelo de IA."
)

doc.add_paragraph(
    "En el prompt, le pedimos a la IA que incluya un campo llamado 'reasoning' "
    "(razonamiento) donde explica paso a paso por que tomo cada decision. "
    "Ademas, le pedimos que incluya 'citas textuales' "
    "(frases exactas del articulo) como evidencia."
)

doc.add_paragraph(
    "Por ejemplo, si la IA decide que una noticia es FALSO, el reasoning podria decir:"
)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(2)
run = p.add_run(
    '"Afirmacion medica extraordinaria sin respaldo de estudio verificable '
    '(el articulo dice: \'descubrimiento que ocultan\'). '
    "Sin autor, fecha ni fuente primaria. Dominio no reconocido.\""
)
run.font.italic = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)

doc.add_paragraph(
    "Esto es como un detective que no solo dice 'es culpable', "
    "sino que te muestra las pruebas: huellas dactilares, testimonios, documentos."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 4. EXTRACCION Y PROCESAMIENTO DEL TEXTO
# ═══════════════════════════════════════════════════════════

doc.add_heading("4. Extraccion y Procesamiento del Texto", level=1)

doc.add_heading("4.1 Flujo completo de datos", level=2)

doc.add_paragraph(
    "Aqui te explico paso a paso como viaja la informacion desde que el usuario "
    "introduce una URL hasta que ve el resultado en pantalla:"
)

steps = [
    ("Paso 1: Input del usuario",
     "El usuario escribe una URL (direccion web) como 'https://milenio.com/noticia' "
     "en el cuadro de texto y hace clic en 'Analizar'. "
     "El frontend crea un AbortController (un temporizador de 60 segundos) "
     "y envia la URL al backend via una peticion POST a /analyze."),

    ("Paso 2: Scraping (leer la pagina)",
     "El backend recibe la URL y empieza a intentar acceder a la pagina web. "
     "Usa 4 estrategias en orden: cloudscraper, curl_cffi, requests, Playwright. "
     "Si todas fallan, intenta con Google Cache (una copia guardada por Google). "
     "El objetivo es obtener el codigo HTML de la pagina."),

    ("Paso 3: Extraccion de texto",
     "Una vez que tiene el codigo HTML, BeautifulSoup lo analiza. "
     "Primero elimina los tags que no sirven (scripts, estilos, menus, pies de pagina). "
     "Luego busca el contenido principal: busca tags <article>, <main>, o <body>. "
     "Extrae todos los parrafos (<p>) que tengan mas de 40 caracteres. "
     "Si no encuentra suficiente texto, busca en todos los parrafos de la pagina."),

    ("Paso 4: Noticias similares",
     "El backend busca en Google News RSS noticias con el mismo titulo "
     "para tener contexto adicional. Es como cuando buscas la misma noticia "
     "en diferentes periodicos para confirmar que es real."),

    ("Paso 5: Construccion del prompt",
     "El backend arma un mensaje (prompt) para Groq que incluye: "
     "el SYSTEM_PROMPT (instrucciones generales, 148 lineas), "
     "el USER_PROMPT_BASE (instrucciones especificas + el texto de la noticia), "
     "3 ejemplos few-shot (ejemplos de como queremos la respuesta), "
     "y el contexto de noticias similares si las hay."),

    ("Paso 6: Llamada a Groq",
     "El backend envia el prompt a Groq API via la libreria groq. "
     "Groq ejecuta el modelo llama-3.3-70b-versatile con temperatura 0.1 "
     "(casi deterministica: siempre da respuestas similares para el mismo texto). "
     "Groq responde con un JSON que contiene el veredicto, la confianza, "
     "el resumen, las afirmaciones, el razonamiento, etc."),

    ("Paso 7: Post-proceso",
     "El backend verifica que la respuesta sea valida. "
     "Si el dominio es confiable (Milenio, Reforma, CNN, etc.) y Groq dice FALSO, "
     "lo cambia automaticamente a NO VERIFICABLE (porque un medio reconocido "
     "no publicaria algo totalmente falso). "
     "Llama a translate_analysis() para agregar campos en ingles si faltan."),

    ("Paso 8: Respuesta al frontend",
     "El backend envia el analisis completo al frontend en formato JSON. "
     "El frontend lo renderiza (muestra en pantalla) con: "
     "colores por veredicto (verde=REAL, rojo=FALSO, etc.), "
     "barra de confianza con 20 segmentos coloreados, "
     "listas de banderas rojas y senales positivas, "
     "noticias similares, y todo en espanol o ingles segun el toggle."),
]

for title, desc in steps:
    p = doc.add_paragraph()
    run = p.add_run(title + ": ")
    run.font.bold = True
    run.font.size = Pt(11)
    p.add_run(desc)

doc.add_heading("4.2 Scraping (como leemos paginas web)", level=2)

doc.add_paragraph(
    "Scraping es como leer un libro y sacar las partes importantes. "
    "Cuando introduces una URL en VERIFEX, el sistema no solo 'mira' la pagina: "
    "la lee completamente, extrae el texto, y lo prepara para que la IA lo analice."
)

doc.add_paragraph(
    "VERIFEX usa 4 estrategias de scraping (llamadas 'fallback strategies'), "
    "ordenadas de mas rapida a mas lenta:"
)

make_table(doc,
    ["Estrategia", "Velocidad", "Que es (explicacion simple)", "Cuando se usa"],
    [
        ["cloudscraper",
         "~2-5 segundos",
         "Una libreria que se hace pasar por un navegador normal "
         "para enganar a Cloudflare (un sistema de seguridad web). "
         "Es como usar un disfraz de delivery para entrar a un edificio.",
         "Primera opcion. Funciona con la mayoria de sitios."],

        ["curl_cffi",
         "~3-7 segundos",
         "Una libreria que imita el comportamiento exacto de navegadores reales "
         "(Chrome, Safari, Firefox). Es como usar un disfraz muy convincente.",
         "Segunda opcion. Se usa cuando cloudscraper no funciona."],

        ["requests",
         "~2-5 segundos",
         "La forma mas basica de hacer peticiones web. "
         "Es como llamar por telefono: haces una pregunta y esperas la respuesta.",
         "Tercera opcion. Funciona con sitios simples sin proteccion."],

        ["Playwright",
         "~15-35 segundos",
         "Un navegador real controlado por computadora. "
         "Abrir Firefox, navegar a la pagina, esperar a que cargue, y extraer el texto. "
         "Es como tener un robot que abre un navegador y lee la pagina por ti.",
         "Ultima opcion. Es el mas lento pero el mas potente."],
    ]
)

doc.add_heading("4.3 Limpieza del texto", level=2)

doc.add_paragraph(
    "Despues de extraer el texto de la pagina web, hay que limpiarlo. "
    "Es como lavar las verduras antes de cocinar: quitas la tierra, las hojas marchitas, "
    "y dejas solo lo que vas a cocinar."
)

doc.add_paragraph("Pasos de limpieza:")

cleanup_steps = [
    "Eliminar tags HTML no deseados: script, style, nav, footer, aside, header, iframe, noscript. "
    "Son como los anuncios y menus que no nos interesan.",

    "Filtrar por longitud: Solo se quedan los parrafos con mas de 40 caracteres. "
    "Los parrafos muy cortos suelen ser botones o titulares, no contenido real.",

    "Limitar cantidad: Se toman maximo 50 parrafos. "
    "Si la pagina tiene 500 parrafos, no necesitamos todos.",

    "Truncamiento: El contenido final se limita a 5000 caracteres "
    "(aproximadamente 1000 palabras). Esto es para que el prompt no sea demasiado largo.",

    "Manejo de paginas especiales: Si es Instagram, extrae el texto del meta tag. "
    "Si es Threads, busca el JSON embebido con los posts. "
    "Si es Facebook, intenta usar la Graph API.",
]

for i, step in enumerate(cleanup_steps, 1):
    doc.add_paragraph(f"{step}", style='List Bullet')

doc.add_heading("4.4 Persistencia (donde guardamos los datos)", level=2)

doc.add_paragraph(
    "VERIFEX NO guarda nada. Es completamente 'stateless' (sin estado). "
    "Esto significa que cada analisis es independiente y no se recuerda entre sesiones."
)

doc.add_paragraph("Esto es intencional por razones de privacidad:")
doc.add_paragraph(
    "HU-16: 'El analisis no almacena ningun dato personal ni historial del usuario'. "
    "Esto esta documentado en las historias de usuario del proyecto.",
    style='List Bullet'
)

doc.add_paragraph(
    "No hay base de datos (no hay MySQL, PostgreSQL, MongoDB, ni nada). "
    "No hay caché persistente (no hay Redis). "
    "No hay usuarios ni autenticacion (no hay login).",
    style='List Bullet'
)

doc.add_paragraph(
    "La unica 'memoria' es la cache en memoria durante la ejecucion "
    "(si dos personas analizan la misma URL al mismo tiempo, el segundo podria aprovechar "
    "el resultado del primero, pero esto es accidental, no intencional).",
    style='List Bullet'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 5. COMUNICACION FRONTEND-BACKEND
# ═══════════════════════════════════════════════════════════

doc.add_heading("5. Comunicacion Frontend-Backend y Despliegue", level=1)

doc.add_heading("5.1 Como viaja la informacion", level=2)

doc.add_paragraph(
    "La comunicacion entre el frontend (lo que ve el usuario) y el backend (el servidor) "
    "se hace via API REST con JSON. Vamos a explicar que significa eso:"
)

doc.add_paragraph(
    'API: Es como un mesero en un restaurante. Tu (el frontend) le dices al mesero (la API) '
    "que quieres, el mesero va a la cocina (el backend), trae tu pedido, y te lo sirve.",
    style='List Bullet'
)

doc.add_paragraph(
    "REST: Es una forma estandar de hacer las peticiones. "
    "Usa verbos como GET (obtener), POST (crear), PUT (actualizar), DELETE (borrar). "
    "VERIFEX solo usa POST (para analizar URLs) y GET (para verificar que el servidor funciona).",
    style='List Bullet'
)

doc.add_paragraph(
    'JSON: Es un formato de texto estructurado, como una tabla pero en texto plano. '
    'Se ve algo asi: {"nombre": "Marco", "edad": 22}. '
    "Es facil de leer para humanos y para computadoras.",
    style='List Bullet'
)

doc.add_paragraph("Flujo de una peticion tipica:")

flow_steps = [
    "1. El usuario escribe 'https://milenio.com/noticia' y hace clic en Analizar",
    "2. El frontend crea un POST a /analyze con body: {\"url\": \"https://milenio.com/noticia\"}",
    "3. Flask recibe la peticion en la funcion analyze()",
    "4. Flask llama a analyze_url(url) en analyzer.py",
    "5. analyzer.py hace scraping, extrae texto, construye prompt, llama a Groq",
    "6. Groq responde con JSON: {\"verdict\": \"REAL\", \"confidence_score\": 85, ...}",
    "7. Flask agrega noticias similares y otros metadatos",
    "8. Flask devuelve el JSON completo al frontend",
    "9. El frontend renderiza los resultados con colores y graficos",
]

for step in flow_steps:
    doc.add_paragraph(step, style='List Bullet')

doc.add_heading("5.2 Formato de la respuesta", level=2)

doc.add_paragraph(
    "El backend devuelve un JSON con esta estructura. Piensa en el JSON como un formulario "
    "lleno con toda la informacion del analisis:"
)

make_table(doc,
    ["Campo", "Tipo", "Que contiene (explicacion simple)"],
    [
        ["analysis", "Objeto JSON", "Todo el analisis de la IA: veredicto, confianza, resumen, afirmaciones, razonamiento, banderas rojas, etc."],
        ["analysis.verdict", "Texto", "El veredicto: REAL, FALSO, SÁTIRA, ESTAFA, o NO VERIFICABLE"],
        ["analysis.confidence_score", "Numero (0-100)", "Que tan segura esta la IA de su decision. 100 = totalmente segura, 0 = no tiene idea"],
        ["analysis.summary", "Texto", "Resumen del articulo en 2-3 oraciones, escrito por la IA"],
        ["analysis.summary_en", "Texto", "Lo mismo pero en ingles"],
        ["analysis.extracted_claims", "Lista de textos", "Las afirmaciones principales del articulo (ej: 'Banxico subio la tasa a 11.25%')"],
        ["analysis.reasoning", "Lista de textos", "Las razones por las que la IA tomo su decision, con citas textuales del articulo"],
        ["analysis.red_flags", "Lista de textos", "Cosas sospechosas que encontro en el articulo (ej: 'Sin autor ni fecha')"],
        ["analysis.positive_signals", "Lista de textos", "Cosas buenas que encontro (ej: 'Medio reconocido', 'Cita fuente oficial')"],
        ["analysis.article_type", "Texto", "Tipo de articulo: informativa, comercial, opinion, clickbait, denuncia"],
        ["analysis.is_scam", "Booleano (true/false)", "Si la IA detecto que es una estafa o fraude"],
        ["similar_news", "Lista de objetos", "Noticias similares encontradas en Google News RSS"],
        ["url_analyzed", "Texto", "La URL que se analizo"],
        ["domain", "Texto", "El dominio del sitio web (ej: milenio.com)"],
        ["is_credible_source", "Booleano", "Si el dominio esta en la lista de 29 medios reconocidos"],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 6. BASE DE DATOS
# ═══════════════════════════════════════════════════════════

doc.add_heading("6. Base de Datos y Persistencia", level=2)

doc.add_paragraph(
    "VERIFEX NO tiene base de datos. Esto es una decision de diseno intencional."
)

doc.add_paragraph(
    "Las razones son:"
)

doc.add_paragraph(
    "Privacidad: No se guarda ningun dato del usuario ni de sus analisis. "
    "Esto cumple con la HU-16 (historia de usuario 16) que dice explicitamente "
    "que no se almacena ningun dato personal.",
    style='List Bullet'
)

doc.add_paragraph(
    "Simplicidad: Sin base de datos, el sistema es mas facil de mantener "
    "y desplegar. No hay que configurar MySQL, PostgreSQL, ni nada.",
    style='List Bullet'
)

doc.add_paragraph(
    "Escalabilidad: Sin estado, el sistema puede escalar facilmente "
    "agregando mas servidores. Cada peticion es independiente.",
    style='List Bullet'
)

doc.add_paragraph(
    "No se almacenan embeddings (vectores numericos que representan el significado del texto) "
    "para futuras busquedas o comparativas. Si en el futuro se quisiera implementar "
    "una funcion de 'historial de analisis', se tendria que agregar una base de datos.",
    style='List Bullet'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 7. TECHNICAL DEBT Y MEJORAS
# ═══════════════════════════════════════════════════════════

doc.add_heading("7. Technical Debt y Mejoras", level=1)

doc.add_paragraph(
    "'Technical Debt' (deuda tecnica) es como una deuda de dinero: "
    "es cuando tomas atajos ahora que te van a costar mas trabajo despues. "
    "Todo proyecto tiene algo de deuda tecnica, y es importante identificarla "
    "para saber donde mejorar."
)

doc.add_heading("7.1 Cuellos de botella", level=2)

doc.add_paragraph(
    "Un 'cuello de botella' es como una zona de trafico: es la parte del sistema "
    "que va mas lenta y frena a todo lo demas."
)

bottlenecks = [
    ("Scraping con Playwright (~35 segundos)",
     "Playwright es como abrir un navegador completo para leer una pagina. "
     "Es muy potente pero muy lento. Actualmente se usa como ultimo recurso, "
     "lo cual es correcto. Pero si muchos usuarios lo activan a la vez, "
     "el servidor puede quedarse sin memoria."),

    ("Groq API como unico proveedor de IA",
     "Si Groq se cae (por mantenimiento, por exceso de uso, o por cualquier razon), "
     "todo VERIFEX deja de funcionar. No hay otro proveedor de IA como respaldo. "
     "Es como tener un solo cajero en un banco: si se enferma, nadie puede hacer transacciones."),

    ("Segunda llamada a Groq para traduccion",
     "La funcion translate_analysis() hace una segunda llamada a Groq "
     "si la primera no devuelve campos en ingles. "
     "Esto duplica el tiempo de respuesta (de 5s a 10s) y el costo de la API."),

    ("Sin cache de resultados",
     "Si dos usuarios analizan la misma URL, se hacen dos llamadas a Groq. "
     "Si hubiera cache, el segundo usuario recibria el resultado instantaneamente."),
]

for title, desc in bottlenecks:
    p = doc.add_paragraph()
    run = p.add_run(title + ": ")
    run.font.bold = True
    p.add_run(desc)

doc.add_heading("7.2 Problemas de seguridad", level=2)

security_issues = [
    ("CORS abierto a todos los origenes",
     "En app.py: CORS(app, origins=\"*\") permite que cualquier sitio web "
     "haga peticiones a VERIFEX. En produccion, deberia restringirse "
     "al dominio de Render (verifex.onrender.com). "
     "Es como dejar la puerta principal de tu casa abierta: cualquiera puede entrar."),

    ("Contenido sin sanitizar",
     "El texto extraido de las paginas web se envia directamente a Groq "
     "sin limpiar caracteres peligrosos. Si una pagina contiene codigo malicioso "
     "(como scripts de phishing), Groq podria procesarlo. "
     "Es como cocinar ingredientes sin lavarlos primero."),

    ("Sin rate limiting (limite de peticiones)",
     "No hay limite de cuantas peticiones puede hacer un usuario. "
     "Un atacante podria hacer miles de peticiones y agotar la cuota de Groq. "
     "Es como un buffet sin limite: alguien podria comer todo la comida."),

    ("API key en variables de entorno (BIEN HECHO)",
     "La clave de Groq esta guardada en server/.env, no en el codigo fuente. "
     "Esto es correcto y seguro. El archivo .env no se sube a GitHub."),
]

for title, desc in security_issues:
    p = doc.add_paragraph()
    run = p.add_run(title + ": ")
    run.font.bold = True
    p.add_run(desc)

doc.add_heading("7.3 Recomendaciones concretas", level=2)

doc.add_paragraph(
    "Aqui van 5 recomendaciones especificas para mejorar VERIFEX, "
    "ordenadas de mayor a menor impacto:"
)

recommendations = [
    ("1. Agregar cache de resultados con Redis",
     "Que es: Redis es una base de datos muy rapida que guarda informacion en memoria. "
     "Es como una pizarra donde escribes las respuestas para no tener que buscarlas cada vez.",
     "Por que: Si alguien analiza 'https://milenio.com/economia', y 5 minutos despues "
     "otro usuario analiza la misma URL, en vez de hacer scraping + Groq otra vez, "
     "se devolveria el resultado guardado. Ahorraria tiempo y dinero.",
     "Como: Guardar en Redis el hash de la URL + el resultado JSON con TTL de 1 hora. "
     "Antes de hacer scraping, verificar si existe en cache."),

    ("2. Fallback a otro proveedor de IA (OpenAI o Anthropic)",
     "Que es: Tener un segundo experto de respaldo por si Groq no esta disponible.",
     "Por que: Si Groq tiene mantenimiento o se cae, VERIFEX seguiria funcionando "
     "con el segundo proveedor. Es como tener un plan B.",
     "Como: Modificar call_groq() para que si Groq falla, intente con OpenAI GPT-4o-mini "
     "(que es mas barato) o Anthropic Claude Haiku."),

    ("3. Rate limiting con Flask-Limiter",
     "Que es: Un sistema que cuenta cuantas peticiones hace cada usuario "
     "y bloquea si excede el limite.",
     "Por que: Evita que un atacante agote la cuota de Groq "
     "y deja que el servicio funcione para todos.",
     "Como: Instalar Flask-Limiter y configurar 10 requests/minuto por direccion IP."),

    ("4. Logging estructurado",
     "Que es: En vez de usar print() para debug, usar el modulo logging de Python "
     "con formato JSON. Es como llevar un libro de contabilidad en vez de "
     "escribir notas sueltas en servilletas.",
     "Por que: Facilita el monitoreo y la deteccion de errores en produccion.",
     "Como: Reemplazar print() por logging.info(), logging.error(), etc. "
     "Y configurar el formato JSON para integracion con herramientas de monitoreo."),

    ("5. Embeddings para noticias similares",
     "Que es: En vez de buscar por titulo en Google News RSS, "
     "usar 'embeddings' (vectores numericos) para encontrar noticias semanticamente similares.",
     "Por que: La busqueda actual por titulo es limitada. "
     "Si la noticia original dice 'Banxico sube tasa' y la similar dice 'Banco de Mexico incrementa intereses', "
     "la busqueda por titulo no la encontraria, pero un embedding si.",
     "Como: Usar sentence-transformers para generar embeddings del titulo original "
     "y buscar en una base de vectores los mas cercanos."),
]

for title, que, por_que, como in recommendations:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    doc.add_paragraph(f"Que es: {que}")
    doc.add_paragraph(f"Por que: {por_que}")
    doc.add_paragraph(f"Como implementarlo: {como}")

# ═══════════════════════════════════════════════════════════
# METRICAS DEL PROYECTO
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
doc.add_heading("Metricas del Proyecto", level=1)

make_table(doc,
    ["Metrica", "Valor", "Que significa"],
    [
        ["Archivos Python", "4", "app.py, analyzer.py, news_finder.py, test_analyzer.py"],
        ["Archivos TypeScript", "3", "App.tsx, App.test.tsx, index.tsx"],
        ["Tests backend", "27", "Pruebas automaticas con pytest"],
        ["Tests frontend", "52", "Pruebas automaticas con vitest"],
        ["Lineas de codigo (backend)", "~2,500", "Todas las funciones de scraping, IA, parsing"],
        ["Lineas de codigo (frontend)", "~800", "Todos los componentes React"],
        ["Dependencias Python", "12", "Librerias en requirements.txt"],
        ["Dependencias Node", "8", "Librerias en package.json"],
        ["Versiones del sistema", "13", "De v0.1.0 a v1.2.0"],
        ["Duracion del proyecto", "~11.5 meses", "Septiembre 2025 - Agosto 2026"],
        ["Dominios confiables", "29", "Medios de comunicacion reconocidos"],
        ["Categorias de veredicto", "5", "REAL, FALSO, SÁTIRA, ESTAFA, NO VERIFICABLE"],
        ["Estrategias de scraping", "4+1", "cloudscraper, curl_cffi, requests, Playwright + Google Cache"],
    ]
)

# ═══════════════════════════════════════════════════════════
# CONCLUSION
# ═══════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_heading("Conclusion", level=1)

doc.add_paragraph(
    "VERIFEX es un proyecto solido y bien estructurado que demuestra "
    "como la Inteligencia Artificial puede aplicarse a un problema real "
    "y relevante: la desinformacion."
)

doc.add_paragraph(
    "Los puntos fuertes del proyecto son:"
)

doc.add_paragraph(
    "Arquitectura limpia: Separacion clara entre frontend y backend, "
    "con una API REST bien definida.",
    style='List Bullet'
)

doc.add_paragraph(
    "Robustez en el scraping: 4 estrategias de fallback garantizan "
    "que el sistema pueda leer la mayoria de las paginas web.",
    style='List Bullet'
)

doc.add_paragraph(
    "Prompt engineering cuidadoso: Instrucciones detalladas, few-shots, "
    "y citas textuales obligatorias mejoran la calidad de las respuestas.",
    style='List Bullet'
)

doc.add_paragraph(
    "Privacidad: El diseño stateless garantiza que no se guardan datos de usuarios.",
    style='List Bullet'
)

doc.add_paragraph(
    "Los puntos a mejorar son:"
)

doc.add_paragraph(
    "Cache de resultados para evitar llamadas duplicadas a Groq.",
    style='List Bullet'
)

doc.add_paragraph(
    "Fallback a otros proveedores de IA para mayor disponibilidad.",
    style='List Bullet'
)

doc.add_paragraph(
    "Rate limiting para proteger contra abuso.",
    style='List Bullet'
)

doc.add_paragraph(
    "En general, VERIFEX es un proyecto academico de alta calidad "
    "que podria evolucionar facilmente a un producto profesional "
    "con las mejoras sugeridas en este informe."
)

# ── Guardar documento ──
output_path = "/Users/maarco_serrano/Downloads/verifex-standalone 2/tesis/Informe_Tecnico_VERIFEX.docx"
doc.save(output_path)
print(f"Documento generado: {output_path}")
