#!/usr/bin/env python3
"""Genera documento .docx con el Glosario de Términos de VERIFEX."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.dirname(os.path.abspath(__file__))

NAVY = "1B3A5C"
DARK_SLATE = "2C3E50"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADING_COLOR = RGBColor(0x1B, 0x3A, 0x5C)
SUBTITLE_GRAY = RGBColor(0x70, 0x70, 0x80)

def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)

def add_term(table, termino, definicion):
    row = table.add_row()
    row.cells[0].text = termino
    row.cells[1].text = definicion
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = "Calibri"

def create_table(doc, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = 1
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(9.5)
            run.font.name = "Calibri"
            run.font.color.rgb = WHITE
        set_cell_shading(cell, DARK_SLATE)
    return table

def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = HEADING_COLOR
    return h

def main():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    # ── Portada ──
    for _ in range(6):
        doc.add_paragraph("")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("VERIFEX")
    run.font.size = Pt(42)
    run.bold = True
    run.font.color.rgb = HEADING_COLOR

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Glosario de Términos")
    run.font.size = Pt(20)
    run.font.color.rgb = SUBTITLE_GRAY

    doc.add_paragraph("")
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub2.add_run("Analizador de Credibilidad de Noticias")
    run.font.size = Pt(13)
    run.font.color.rgb = SUBTITLE_GRAY
    run.italic = True

    doc.add_page_break()

    # ─═ A ═─
    add_heading_styled(doc, "A", 0)
    tbl = create_table(doc, ["Término", "Definición"])
    add_term(tbl, "AbortController",
        "API nativa de JavaScript que permite cancelar una solicitud fetch antes de que "
        "complete. VERIFEX la usa para implementar un timeout de 60s en el análisis. "
        "Si el backend no responde en ese lapso, el frontend aborta la petición y muestra "
        "un mensaje de error al usuario.")
    add_term(tbl, "adjustedVerdict (veredicto ajustado)",
        "Valor derivado en el frontend (useMemo) que modifica el veredicto original de Groq "
        "según la puntuación de confianza: si el score es menor a 50, se fuerza FALSO; "
        "si está entre 50 y 69, se muestra DUDOSO. Esta lógica no aplica para redes sociales.")
    add_term(tbl, "analyze_url()",
        "Función principal del backend (server/analyzer.py) que orquesta todo el flujo de análisis: "
        "verifica la API key, ejecuta scrape_url(), busca noticias similares, construye el prompt "
        "con contexto, llama a Groq, parsea la respuesta y aplica reclassificación si el dominio "
        "es creíble y el veredicto es FALSO.")
    add_term(tbl, "ApiResponse (interface)",
        "Interfaz TypeScript que define la estructura de la respuesta del backend. Contiene: "
        "analysis (Analysis | null), similar_news (NewsItem[]), url_analyzed, article_text, "
        "domain, is_credible_source y error.")
    add_term(tbl, "app.py",
        "Archivo principal del backend Flask (server/app.py). Define las rutas HTTP "
        "(POST /analyze, GET /health), configura CORS, carga variables de entorno con "
        "python-dotenv y sirve el frontend compilado como SPA.")
    add_term(tbl, "article_type (tipo de artículo)",
        "Campo en la respuesta JSON de Groq que clasifica la naturaleza del contenido: "
        "'informativa' (noticia neutral), 'comercial' (promoción), 'opinion' (columna/editorial), "
        "'clickbait' (titular engañoso) o 'denuncia' (investigación).")
    add_term(tbl, "article_text (texto del artículo)",
        "Texto extraído del HTML de la URL analizada, limitado a 2000 caracteres. "
        "Se muestra al usuario como vista previa en el frontend.")
    add_term(tbl, "article_type (tipo de artículo)",
        "Clasificación del contenido según su naturaleza. Groq la determina durante el análisis "
        "y se muestra como un badge en la UI. Los tipos son: informativa, comercial, opinion, "
        "clickbait y denuncia.")
    doc.add_paragraph("")

    # ─═ B ═─
    add_heading_styled(doc, "B", 0)
    tbl = create_table(doc, ["Término", "Definición"])
    add_term(tbl, "Backend",
        "Servidor Python/Flask que procesa las solicitudes de análisis. Corre en localhost:5001 "
        "en desarrollo y en Render en producción. Expone endpoints /analyze y /health.")
    add_term(tbl, "banderas rojas (red_flags)",
        "Lista de alertas detectadas por Groq durante el análisis. Incluyen: titular sensacionalista, "
        "falta de autor o fuentes, lenguaje alarmista, contradicción titular-contenido, "
        "promesas de curas milagrosas, entre otras. Se muestran en el componente RedFlags.")
    add_term(tbl, "BeautifulSoup (bs4)",
        "Librería de Python para parsear HTML y XML. VERIFEX la usa en _extract_from_html() "
        "para navegar el DOM: encuentra el título, meta descripción, párrafos, elimina etiquetas "
        "no deseadas (script, style, nav, footer) y extrae el texto del artículo.")
    add_term(tbl, "build.sh",
        "Script de shell (build.sh en la raíz) que Render ejecuta durante el build. "
        "Instala dependencias Python y los navegadores Firefox/Chromium de Playwright. "
        "Es crítico para que el scraper funcione en producción.")
    add_term(tbl, "BROWSER_HEADERS",
        "Diccionario de headers HTTP en analyzer.py que imita una solicitud de navegador real: "
        "Accept, Accept-Language (es-MX), Sec-Fetch-*, Cache-Control. Se usa en todas las "
        "estrategias de scraping para evitar bloqueos por User-Agent o fingerprint.")
    doc.add_paragraph("")

    # ─═ C ═─
    add_heading_styled(doc, "C", 0)
    tbl = create_table(doc, ["Término", "Definición"])
    add_term(tbl, "call_groq()",
        "Función que envía los prompts a la API de Groq. Primero intenta con el modelo "
        "llama-3.3-70b-versatile; si falla, usa llama-3.1-8b-instant como fallback. "
        "Configura temperature=0.1 y response_format='json_object' para respuestas "
        "consistentes y parseables.")
    add_term(tbl, "cita textual",
        "Requerimiento del prompt de sistema que obliga a Groq a incluir frases exactas "
        "del artículo entre comillas en cada punto de reasoning y extracted_claims. "
        "Garantiza que el análisis esté fundamentado en el contenido real y no sea inventado.")
    add_term(tbl, "clickbait",
        "Tipo de artículo (article_type) con titular engañoso que no refleja el contenido real. "
        "Por ejemplo, un titular que dice 'IMPACTANTE: esto cambiará tu vida' pero el contenido "
        "es trivial. Se marca como bandera roja.")
    add_term(tbl, "cloudscraper",
        "Librería de Python que evade Cloudflare imitando el handshake TLS de un navegador. "
        "Es la primera estrategia de scraping. Configura 4 perfiles (Chrome/Firefox en Linux/Windows) "
        "con 2 intentos cada uno (8 totales). Timeout de 30s por intento.")
    add_term(tbl, "Cloudflare",
        "Servicio de protección web que muchas páginas usan. Presenta un challenge JavaScript "
        "que bloquea solicitudes automatizadas. VERIFEX implementa 5 estrategias para evadirlo: "
        "cloudscraper, curl_cffi, requests, Playwright y Google Cache.")
    add_term(tbl, "ConfidenceBar (componente)",
        "Componente React (ConfidenceBar.tsx) que renderiza una barra de 20 segmentos iluminados. "
        "Color según rango: rojo (0-40), naranja (41-69), cian (70-89), verde (90-100). "
        "Representa visualmente el nivel de confianza del análisis.")
    add_term(tbl, "confidence_score",
        "Puntuación de confianza del 0 al 100 que Groq asigna al veredicto. Se muestra en "
        "la ConfidenceBar y en el frontend se usa para ajustar el veredicto: <50 → FALSO, "
        "50-69 → DUDOSO, ≥70 → se respeta el original.")
    add_term(tbl, "CORS (Cross-Origin Resource Sharing)",
        "Mecanismo de seguridad del navegador que permite solicitudes entre diferentes orígenes. "
        "VERIFEX configura Flask-CORS con origins='*' para permitir que el frontend en "
        "localhost:5173 se comunique con el backend en localhost:5001.")
    add_term(tbl, "CREDIBLE_DOMAINS",
        "Conjunto de 31 dominios de medios de comunicación reconocidos (jornada.com.mx, "
        "reuters.com, eluniversal.com.mx, etc.). Si el dominio de la URL está en este conjunto "
        "y Groq retorna FALSO, se reclasifica automáticamente a NO VERIFICABLE para evitar "
        "falsos positivos en fuentes confiables.")
    add_term(tbl, "curl_cffi",
        "Segunda estrategia de scraping. Usa la librería curl_cffi para impersonar versiones "
        "específicas de navegador (chrome123, chrome120, safari17_0, chrome124) imitando el "
        "TLS fingerprint exacto. Timeout de 30s por intento.")
    doc.add_paragraph("")

    # ─═ D ═─
    add_heading_styled(doc, "D", 0)
    tbl = create_table(doc, ["Término", "Definición"])
    add_term(tbl, "despliegue (deployment)",
        "Proceso de poner el sistema en producción. VERIFEX usa Render: el backend como "
        "Web Service con Gunicorn y el frontend como Static Site con build npm. "
        "Incluye build automático, SSL y dominio público.")
    add_term(tbl, "DUDOSO",
        "Veredicto ajustado que solo existe en el frontend. Cuando el confidence_score está "
        "entre 50 y 69, el frontend modifica el veredicto original a DUDOSO. No es un veredicto "
        "de Groq — es una decisión visual del frontend para casos de confianza media-baja.")
    add_term(tbl, "dominio (domain)",
        "Parte de la URL que identifica el sitio web (ej: 'jornada.com.mx'). Se extrae con "
        "get_domain() y se usa para: verificar si es fuente creíble, determinar el prompt "
        "de redes sociales, y mostrar al usuario.")
    doc.add_paragraph("")

    # ─═ E ═─
    add_heading_styled(doc, "E", 0)
    tbl = create_table(doc, ["Término", "Definición"])
    add_term(tbl, "edge cases (casos límite)",
        "Situaciones excepcionales que el sistema debe manejar: URLs inválidas, timeouts, "
        "páginas bloqueadas, contenido vacío, redes sociales sin API, dominios desconocidos, "
        "JSON mal formado de Groq, entre otros. Cada caso tiene un mensaje de error específico.")
    add_term(tbl, "ESTAFA",
        "Veredicto para contenido diseñado para defraudar: phishing, productos milagro, "
        "inversiones falsas, suplantación de identidad. Se diferencia de FALSO porque "
        "implica intención de lucro o daño económico.")
    add_term(tbl, "_extract_from_html()",
        "Función que parsea el HTML obtenido por las estrategias de scraping. Extrae título, "
        "meta descripción, párrafos relevantes (>40 caracteres), maneja casos especiales "
        "(Instagram, Google Cache, Threads) y retorna {content, title, article_text}.")
    add_term(tbl, "extracted_claims (afirmaciones extraídas)",
        "Lista de afirmaciones principales que Groq extrae del artículo, cada una con cita "
        "textual. Se muestran en el frontend como una lista con viñetas.")
    add_term(tbl, "_extract_facebook_post_id()",
        "Función auxiliar que extrae el ID de un post de Facebook desde la URL. "
        "Soporta formatos: /posts/{id}, /videos/{id}, /photos/{id}, story.php?story_fbid=, "
        "photo.php?fbid=.")
    doc.add_paragraph("")

    # ─═ F ═─
    add_heading_styled(doc, "F", 0)
    tbl = create_table(doc, ["Término", "Definición"])
    add_term(tbl, "Facebook Graph API",
        "API oficial de Facebook para acceder a contenido público. VERIFEX la usa como fallback "
        "cuando el scraper encuentra una página de inicio de sesión. Requiere FACEBOOK_APP_ID "
        "y FACEBOOK_APP_SECRET configurados en .env.")
    add_term(tbl, "FALSO",
        "Veredicto que indica que el contenido contiene información demostrablemente incorrecta: "
        "afirmaciones que contradicen hechos establecidos, citas fabricadas, estadísticas "
        "inventadas o teorías conspirativas sin respaldo. NO debe usarse por 'no poder verificar'.")
    add_term(tbl, "fallback",
        "Mecanismo de respaldo en cascada. Aparece en múltiples niveles: 5 estrategias de "
        "scraping (si una falla, prueba la siguiente), 2 modelos de Groq (si llama-3.3-70b "
        "falla, usa llama-3.1-8b), y en el análisis de Facebook (scraper → Graph API).")
    add_term(tbl, "feedparser",
        "Librería de Python para parsear feeds RSS/Atom. VERIFEX la usa para procesar "
        "las respuestas de Google News RSS y extraer título, URL, fecha y fuente de cada noticia.")
    add_term(tbl, "Fetch API",
        "API nativa del navegador para hacer solicitudes HTTP. VERIFEX la usa para enviar "
        "el POST /analyze con la URL. Se combinó con AbortController para manejar timeouts "
        "de 60s. No se usó Axios para reducir el bundle size.")
    add_term(tbl, "few-shot examples (ejemplos contextuales)",
        "Técnica de prompt engineering donde se incluyen ejemplos completos en el prompt "
        "para guiar al LLM. VERIFEX proporciona 3 ejemplos: REAL (noticia periodística), "
        "FALSO (afirmación médica falsa) y NO VERIFICABLE (testimonio ambiguo), cada uno "
        "con el JSON de respuesta esperado.")
    add_term(tbl, "Flask",
        "Framework web de Python usado para el backend. Ligero, con solo 2 endpoints. "
        "Se eligió sobre Django por su simplicidad y rápida inicialización.")
    add_term(tbl, "FLASK_PORT",
        "Variable de entorno que define el puerto del servidor Flask. Por defecto es 5001. "
        "Se configura en server/.env y se lee en app.py con os.getenv().")
    add_term(tbl, "frontend",
        "Interfaz de usuario construida con React 18 + TypeScript + Vite + Tailwind CSS. "
        "SPA de una sola vista que se comunica con el backend vía HTTP. Incluye 6 componentes "
        "y carga diferida (React.lazy) para SimilarNews.")
    doc.add_paragraph("")

    # ─═ G ═─
    add_heading_styled(doc, "G", 0)
    tbl = create_table(doc, ["Término", "Definición"])
    add_term(tbl, "get_domain()",
        "Función utilitaria que extrae el dominio de una URL usando urllib.parse.urlparse "
        "y elimina el prefijo 'www.' si existe.")
    add_term(tbl, "get_groq_client()",
        "Función que obtiene el cliente de Groq usando la API key de GROQ_API_KEY. "
        "Retorna None si no hay API key, lo que causa que el sistema devuelva error 503.")
    add_term(tbl, "glitch effect",
        "Efecto visual cyberpunk aplicado al título 'VERIFEX' en el frontend. "
        "Consiste en un pseudo-elemento con datos distorsionados y animación CSS "
        "que simula fallas de video.")
    add_term(tbl, "Google Cache (Google Web Cache)",
        "Quinta y última estrategia de scraping. Obtiene el artículo desde el caché de Google "
        "(webcache.googleusercontent.com). Google ya descargó la página con su IP, "
        "por lo que evade completamente Cloudflare. Timeout de 20s por intento.")
    add_term(tbl, "Google News RSS",
        "Feed RSS de Google News que VERIFEX usa para buscar noticias relacionadas. "
        "Configurado con hl=es&gl=MX para resultados en español de México. "
        "Es gratuito, sin autenticación, y devuelve título, URL, fecha y fuente.")
    add_term(tbl, "Groq API",
        "API de inferencia de LLMs ultrarrápida. VERIFEX la usa para clasificar la credibilidad "
        "de noticias. Modelo principal: llama-3.3-70b-versatile; fallback: llama-3.1-8b-instant. "
        "Se eligió sobre OpenAI por menor latencia y costo.")
    add_term(tbl, "GROQ_API_KEY",
        "Variable de entorno obligatoria que contiene la clave de API de Groq. Sin ella, "
        "el backend devuelve error 503. Se configura en server/.env. Obtener en console.groq.com.")
    add_term(tbl, "Gunicorn",
        "Servidor WSGI para Python en producción. VERIFEX lo usa en Render con "
        "2 workers y timeout de 120s. Comando: gunicorn app:app --chdir server "
        "--bind 0.0.0.0:$PORT --timeout 120 --workers 2.")
    doc.add_paragraph("")

    # ─═ H ═─
    add_heading_styled(doc, "H", 0)
    tbl = create_table(doc, ["Término", "Definición"])
    add_term(tbl, "/health",
        "Endpoint GET que retorna {'status': 'ok'}. Usado por Render para health checks "
        "y monitoreo básico del servidor.")
    add_term(tbl, "HTML parsing",
        "Proceso de analizar el HTML de una página web para extraer información estructurada. "
        "VERIFEX usa BeautifulSoup con lxml para extraer título, meta descripción y párrafos.")
    add_term(tbl, "HTTP_PROXY / HTTPS_PROXY",
        "Variables de entorno para configurar un proxy opcional. Útil cuando Render bloquea "
        "sitios por IP de datacenter. Se aplican en todas las estrategias de scraping mediante "
        "SCRAPING_PROXY.")
    add_term(tbl, "_http_get() (orquestador de scraping)",
        "Función que ejecuta las 5 estrategias de scraping en secuencia: cloudscraper → "
        "curl_cffi → requests → playwright → google_cache. En cuanto una retorna éxito, "
        "la retorna inmediatamente. Si todas fallan, concatena los errores.")
    add_term(tbl, "HU (Historia de Usuario / User Story)",
        "Requerimiento funcional expresado desde la perspectiva del usuario. Formato: "
        "'[HU-NN] Descripción'. VERIFEX tiene 34 HUs (HU-01 a HU-34) que cubren desde "
        "pegar una URL hasta pruebas automatizadas.")
    doc.add_paragraph("")

    # ─═ I ═─
    add_heading_styled(doc, "I", 0)
    tbl = create_table(doc, ["Término", "Definición"])
    add_term(tbl, "is_credible_source",
        "Booleano que indica si el dominio de la URL está en CREDIBLE_DOMAINS. "
        "Se muestra en el frontend como un badge verde 'Fuente Verificada'.")
    add_term(tbl, "is_scam",
        "Booleano en la respuesta de Groq que indica si el contenido es una estafa. "
        "Se muestra en el frontend como una alerta roja pulsante 'ALERTA DE ESTAFA'.")
    add_term(tbl, "Instagram scraping",
        "Caso especial en _extract_from_html(). Para URLs de Instagram, usa la meta description "
        "como texto principal y elimina comentarios del DOM.")
    add_term(tbl, "integration testing (pruebas de integración)",
        "Pruebas que verifican la comunicación entre frontend, backend y Groq API. "
        "VERIFEX incluye 27 tests de backend con pytest y 52 tests de frontend con Vitest.")
    doc.add_paragraph("")

    # ─═ J ═─
    add_heading_styled(doc, "J", 0)
    tbl = create_table(doc, ["Término", "Definición"])
    add_term(tbl, "JSON (JavaScript Object Notation)",
        "Formato de intercambio de datos. VERIFEX lo usa para: la solicitud POST /analyze "
        "({url: string}), la respuesta de Groq (response_format='json_object'), y la respuesta "
        "del backend (ApiResponse).")
    add_term(tbl, "jsdom",
        "Entorno DOM simulado para pruebas. VERIFEX lo usa en vitest.config.ts para ejecutar "
        "pruebas de componentes React sin navegador real.")
    doc.add_paragraph("")

    # ─═ K ═─
    add_heading_styled(doc, "K", 0)
    tbl = create_table(doc, ["Término", "Definición"])
    add_term(tbl, "Kanban",
        "Metodología visual de gestión de proyectos. VERIFEX documenta su progreso en un "
        "tablero Kanban con columnas: Backlog, Stories, Por Hacer, En Proceso, Por Verificar "
        "y Acabado/Terminado. El tablero se genera automáticamente con openpyxl.")
    doc.add_paragraph("")

    # ─═ L ═─
    add_heading_styled(doc, "L", 0)
    tbl = create_table(doc, ["Término", "Definición"])
    add_term(tbl, "Lang (type alias)",
        "Tipo unión TypeScript: 'es' | 'en'. Representa los idiomas soportados por la interfaz. "
        "Controla las traducciones (TRANSLATIONS) y se pasa a todos los componentes hijos.")
    add_term(tbl, "LanguageToggle (componente)",
        "Botón que alterna entre español e inglés. Muestra el código del idioma opuesto. "
        "Usa el callback handleToggleLang del componente App.")
    add_term(tbl, "llama-3.3-70b-versatile",
        "Modelo principal de Groq para clasificación. 70B parámetros, versátil para "
        "tareas de análisis. Si falla, VERIFEX usa llama-3.1-8b-instant como fallback.")
    add_term(tbl, "llama-3.1-8b-instant",
        "Modelo de respaldo de Groq. 8B parámetros, más rápido pero menos preciso. "
        "Se usa cuando llama-3.3-70b-versatile no está disponible.")
    add_term(tbl, "LOGIN_PATTERNS",
        "Lista de patrones de texto en español e inglés que indican una página de inicio "
        "de sesión: 'iniciar sesión', 'contraseña', 'sign up to see', etc. Si se detectan "
        "2 o más coincidencias, el scraper asume que la página está bloqueada.")
    add_term(tbl, "lxml",
        "Parser de XML/HTML para Python. VERIFEX lo usa como parser de BeautifulSoup "
        "por ser más rápido que el parser nativo de Python.")
    doc.add_paragraph("")

    # ─═ M ═─
    add_heading_styled(doc, "M", 0)
    tbl = create_table(doc, ["Término", "Definición"])
    add_term(tbl, "models / modelos de Groq",
        "VERIFEX usa dos modelos de Groq en orden de preferencia: "
        "1. llama-3.3-70b-versatile (principal, 70B params)\n"
        "2. llama-3.1-8b-instant (fallback, 8B params)")
    add_term(tbl, "multi-strategy scraping",
        "Arquitectura de 5 estrategias de scraping en cascada: cloudscraper, curl_cffi, "
        "requests, Playwright y Google Cache. Cada una usa técnicas diferentes para "
        "maximizar la probabilidad de extraer contenido exitosamente.")
    doc.add_paragraph("")

    # ─═ N ═─
    add_heading_styled(doc, "N", 0)
    tbl = create_table(doc, ["Término", "Definición"])
    add_term(tbl, "news_finder.py",
        "Módulo backend que busca noticias relacionadas vía Google News RSS. "
        "Contiene la función find_similar_news(query, max_results) que retorna hasta "
        "5 resultados con título, URL, fecha y fuente.")
    add_term(tbl, "NO VERIFICABLE",
        "Veredicto por defecto para casos dudosos. Se aplica cuando el contenido hace "
        "afirmaciones serias pero sin fuentes verificables, o hay ambigüedad sin señales "
        "claras de fabricación. También se usa como reclasificación cuando un dominio "
        "creíble recibe FALSO de Groq.")
    add_term(tbl, "npm / npm install",
        "Gestor de paquetes de Node.js. Se usa para instalar las dependencias del frontend "
        "(React, TypeScript, Vite, Tailwind, Testing Library). Comando: npm install.")
    doc.add_paragraph("")

    # ─═ O ═─
    add_heading_styled(doc, "O", 0)
    tbl = create_table(doc, ["Término", "Definición"])
    add_term(tbl, "openpyxl",
        "Librería de Python para leer/escribir archivos Excel. VERIFEX la usa para generar "
        "el tablero Kanban (Kanban_VERIFEX.xlsx) con estilos, colores y formato profesional.")
    doc.add_paragraph("")

    # ─═ P ═─
    add_heading_styled(doc, "P", 0)
    tbl = create_table(doc, ["Término", "Definición"])
    add_term(tbl, "parse_response()",
        "Función que parsea la respuesta JSON de Groq. Maneja casos donde el JSON viene "
        "envuelto en bloques ```, y como fallback busca el primer '{' y último '}' "
        "en el texto para extraer el JSON.")
    add_term(tbl, "Playwright",
        "Cuarta estrategia de scraping. Lanza un navegador real (Firefox o Chromium) "
        "headless con perfil completo: viewport 1920x1080, locale es-MX, timezone "
        "America/Mexico_City, --no-sandbox. Espera hasta 35s a que Cloudflare resuelva "
        "el challenge. Timeout total por intento: ~95s.")
    add_term(tbl, "Polling (sondeo periódico)",
        "Técnica usada en Playwright para esperar a que Cloudflare resuelva el challenge. "
        "Revisa cada 1 segundo (hasta 35 veces) si el contenido ya no contiene indicadores "
        "de Cloudflare o si el título tiene contenido real.")
    add_term(tbl, "positive_signals (señales positivas)",
        "Lista de indicadores de credibilidad detectados por Groq: medio reconocido, "
        "cita fuente oficial, reporta hecho verificable, lenguaje neutral, "
        "autor identificable, entre otros.")
    add_term(tbl, "Procfile",
        "Archivo de configuración de Render que define el comando de inicio del Web Service. "
        "VERIFEX usa: web: python -m gunicorn app:app --chdir server --bind 0.0.0.0:$PORT "
        "--timeout 120 --workers 2 --log-level info")
    add_term(tbl, "prompt engineering",
        "Diseño de instrucciones para el LLM. VERIFEX implementa: system prompt con reglas "
        "de clasificación, user prompt con el contenido y contexto, few-shot examples, "
        "cita textual obligatoria, prompt especial para redes sociales, y response_format "
        "forzado a json_object.")
    add_term(tbl, "proxy (HTTP_PROXY)",
        "Servidor intermediario opcional para scraping. Configurable vía HTTP_PROXY o "
        "HTTPS_PROXY del entorno. Útil para evadir bloqueos por IP de datacenter en Render.")
    add_term(tbl, "pytest",
        "Framework de pruebas para Python. VERIFEX tiene 27 tests en server/test_analyzer.py "
        "que cubren: extracción de dominio, parseo de JSON, scraping, y análisis completo.")
    add_term(tbl, "python-dotenv",
        "Librería para cargar variables de entorno desde archivo .env. VERIFEX la usa "
        "en app.py con load_dotenv() para leer GROQ_API_KEY, FLASK_PORT, etc.")
    doc.add_paragraph("")

    # ─═ Q ═─
    add_heading_styled(doc, "Q", 0)
    tbl = create_table(doc, ["Término", "Definición"])
    add_term(tbl, "query string",
        "Parámetros en la URL después del signo '?'. Google News RSS los usa: "
        "q={query}&hl=es&gl=MX. También se parsean en URLs de Facebook para extraer IDs.")
    doc.add_paragraph("")

    # ─═ R ═─
    add_heading_styled(doc, "R", 0)
    tbl = create_table(doc, ["Término", "Definición"])
    add_term(tbl, "React 18",
        "Librería frontend para construir interfaces de usuario. VERIFEX usa componentes "
        "funcionales con hooks (useState, useCallback, useMemo, lazy, Suspense). No usa "
        "clases ni Redux.")
    add_term(tbl, "React.lazy + Suspense",
        "Técnica de code splitting que carga el componente SimilarNews solo cuando es "
        "necesario (cuando hay noticias similares). Reduce el bundle inicial ~15%.")
    add_term(tbl, "REAL",
        "Veredicto para contenido que proviene de un medio establecido, reporta hechos "
        "de manera periodística estándar, cita fuentes identificables y usa lenguaje neutral. "
        "Un artículo normal de Milenio, Reuters, BBC o similar se clasifica como REAL "
        "aunque no se pueda verificar cada detalle individualmente.")
    add_term(tbl, "reasoning (razonamiento)",
        "Lista de razones detalladas que justifican el veredicto, cada una con cita textual "
        "del artículo. Se muestra en el frontend enumerada con formato '01', '02', etc.")
    add_term(tbl, "reclasificación / override",
        "Lógica que modifica el veredicto de Groq: si el dominio está en CREDIBLE_DOMAINS "
        "y Groq retorna FALSO, se reclasifica a NO VERIFICABLE. Previene falsos positivos "
        "en medios de comunicación reconocidos.")
    add_term(tbl, "RedFlags (componente)",
        "Componente React que muestra dos paneles: uno rojo con banderas rojas detectadas "
        "y otro cian con señales positivas. Cada ítem se renderiza como un pill con estilo cyberpunk.")
    add_term(tbl, "Render",
        "Plataforma de despliegue en la nube. VERIFEX despliega: backend como Web Service "
        "(Gunicorn + Flask) y frontend como Static Site (npm run build). Render ofrece "
        "build automático, SSL gratuito, dominio público y tier free.")
    add_term(tbl, "requests (librería)",
        "Tercera estrategia de scraping. La librería HTTP estándar de Python. 2 intentos: "
        "primero con verify=True, segundo con verify=False. Timeout de 15s por intento. "
        "Es la más rápida pero la más fácil de bloquear.")
    add_term(tbl, "requirements.txt",
        "Archivo con dependencias de Python. VERIFEX tiene dos: server/requirements.txt "
        "(flask, flask-cors, requests, cloudscraper, curl_cffi, playwright, beautifulsoup4, "
        "python-dotenv, lxml, groq, gunicorn) y uno en la raíz para detección de Render.")
    add_term(tbl, "response_format='json_object'",
        "Parámetro de la API de Groq que fuerza al modelo a responder en JSON válido. "
        "Crítico para que parse_response() pueda procesar la respuesta consistentemente.")
    add_term(tbl, "RSS (Really Simple Syndication)",
        "Formato XML para sindicación de contenido. VERIFEX consume Google News RSS "
        "para obtener noticias relacionadas. Lo parsea con xml.etree.ElementTree.")
    doc.add_paragraph("")

    # ─═ S ═─
    add_heading_styled(doc, "S", 0)
    tbl = create_table(doc, ["Término", "Definición"])
    add_term(tbl, "SÁTIRA",
        "Veredicto para contenido de humor o parodia. Solo se aplica si el formato, "
        "tono y contexto indican claramente sátira. En redes sociales, se fuerza confianza >=80.")
    add_term(tbl, "scraper / scraping",
        "Proceso de extracción automatizada de contenido web. VERIFEX implementa 5 estrategias "
        "en cascada para maximizar la tasa de éxito, desde cloudscraper hasta Google Cache.")
    add_term(tbl, "scrape_url()",
        "Función que orquesta el scraping completo de una URL. Llama a _http_get(), "
        "extrae contenido con _extract_from_html(), reintenta con Playwright si el texto "
        "es muy corto, detecta login blocking y maneja Facebook con Graph API.")
    add_term(tbl, "SCRAPING_PROXY",
        "Variable de entorno que almacena la URL del proxy para scraping. Se lee de "
        "HTTP_PROXY o HTTPS_PROXY del entorno. Aplica a todas las estrategias.")
    add_term(tbl, "SimilarNews (componente)",
        "Componente React con carga diferida (React.lazy) que muestra una cuadrícula "
        "de tarjetas con noticias relacionadas. Cada tarjeta muestra fuente, título y fecha.")
    add_term(tbl, "SPA (Single Page Application)",
        "Arquitectura frontend donde toda la aplicación se carga en una sola página HTML "
        "y React Router maneja la navegación. VERIFEX usa una SPA de una sola vista.")
    add_term(tbl, "summary (resumen)",
        "Resumen neutral del artículo en 2-3 oraciones generado por Groq. Se muestra en "
        "un panel debajo del veredicto y la barra de confianza.")
    add_term(tbl, "SYSTEM_PROMPT",
        "Prompt del sistema para Groq. Contiene las reglas de clasificación, banderas rojas "
        "a detectar, tipos de artículo, instrucciones de cita textual y el formato exacto "
        "de la respuesta JSON.")
    add_term(tbl, "Tailwind CSS",
        "Framework CSS utilitario. VERIFEX lo usa para estilos responsive, clases de "
        "grid, padding, colores y efectos. Se configuró con PostCSS y autoprefixer.")
    add_term(tbl, "temperature",
        "Parámetro del LLM que controla la aleatoriedad de las respuestas. "
        "VERIFEX usa temperature=0.1 (muy baja) para respuestas consistentes y "
        "determinísticas. A 0, el modelo siempre da la misma respuesta para el mismo input.")
    add_term(tbl, "test_analyzer.py",
        "Archivo de pruebas del backend (server/test_analyzer.py). Contiene 27 tests "
        "de pytest que cubren: dominio, parseo de JSON, scraping y análisis.")
    add_term(tbl, "Testing Library (@testing-library/react)",
        "Librería para pruebas de componentes React. VERIFEX la usa con Vitest para "
        "52 tests distribuidos en 7 archivos (UrlInput, VerdictDisplay, ConfidenceBar, "
        "RedFlags, SimilarNews, LanguageToggle).")
    add_term(tbl, "Threads",
        "Red social de Instagram (threads.net / threads.com). VERIFEX scraping extrae "
        "posts del JSON embebido en <script type='application/json'> buscando "
        "'text_post_app_thread'.")
    add_term(tbl, "timeout",
        "Límite de tiempo para operaciones. VERIFEX maneja: 60s para la solicitud frontend "
        "(AbortController), 30s para cloudscraper/curl_cffi, 15s para requests, ~95s para "
        "Playwright, 20s para Google Cache, 10s para Google News RSS.")
    add_term(tbl, "TLS fingerprint (huella TLS)",
        "Identificador único del handshake TLS de un navegador. curl_cffi lo impersona "
        "para parecer un navegador real y evitar bloqueos por fingerprinting.")
    add_term(tbl, "TRANSLATIONS",
        "Objeto TypeScript con traducciones español/inglés para toda la interfaz. "
        "Contiene ~20 claves por idioma: subtitle, summary, claims, reasoning, "
        "errorTitle, urlAnalyzed, etc. Controlado por el estado lang.")
    add_term(tbl, "_try_cloudscraper()",
        "Primera estrategia de scraping. Usa cloudscraper con 4 perfiles de navegador "
        "y 2 intentos cada uno. Retorna (respuesta, error).")
    add_term(tbl, "_try_curl_cffi()",
        "Segunda estrategia de scraping. Usa curl_cffi impersonando chrome123, chrome120, "
        "safari17_0, chrome124. Retorna (respuesta, error).")
    add_term(tbl, "_try_facebook_graph_api()",
        "Función que obtiene contenido de Facebook vía Graph API. Requiere credenciales "
        "en .env. Retorna dict con content, title, article_text o None si falla.")
    add_term(tbl, "_try_google_cache()",
        "Quinta estrategia de scraping. Usa webcache.googleusercontent.com con 2 intentos. "
        "Retorna (respuesta, error).")
    add_term(tbl, "_try_playwright()",
        "Cuarta estrategia de scraping. Lanza Firefox/Chromium con Playwright. "
        "2 intentos por motor, espera hasta 35s a que Cloudflare resuelva. "
        "Retorna (respuesta, error).")
    add_term(tbl, "_try_requests()",
        "Tercera estrategia de scraping. requests estándar con 2 intentos "
        "(verify=True, verify=False). Timeout 15s. Retorna (respuesta, error).")
    add_term(tbl, "TypeScript",
        "Superconjunto de JavaScript con tipado estático. VERIFEX lo usa en todo el "
        "frontend: define interfaces (Analysis, NewsItem, ApiResponse), tipos (Lang), "
        "y proporciona type checking en compilación.")
    doc.add_paragraph("")

    # ─═ U ═─
    add_heading_styled(doc, "U", 0)
    tbl = create_table(doc, ["Término", "Definición"])
    add_term(tbl, "UML (Unified Modeling Language)",
        "Lenguaje de modelado visual. VERIFEX documenta su arquitectura con diagramas de: "
        "casos de uso, actividades, secuencia, clases y entidad-relación.")
    add_term(tbl, "UrlInput (componente)",
        "Componente React con input de URL y botón de análisis. Maneja su estado local (url), "
        "validación (no vacío, no loading), y callbacks handleSubmit, handleClear. "
        "Incluye subcomponente LoadingSegments con animación de barras.")
    add_term(tbl, "USER_AGENTS",
        "Lista de 5 User-Agent rotativos para evitar bloqueos por fingerprint: Chrome "
        "macOS/Windows/Linux, Safari macOS, Firefox Windows. Se selecciona uno aleatorio "
        "en cada intento de scraping.")
    add_term(tbl, "USER_PROMPT_BASE",
        "Template del prompt de usuario para Groq. Se llena dinámicamente con: URL, dominio, "
        "indicador de credibilidad, contenido extraído, contexto de otras fuentes. "
        "Define el formato exacto del JSON de respuesta esperado.")
    add_term(tbl, "useCallback / useMemo / useState",
        "Hooks de React. useCallback memoriza funciones (handleAnalyze, handleToggleLang). "
        "useMemo memoriza valores calculados (adjustedVerdict). useState maneja estado local "
        "(lang, loading, result, error).")
    doc.add_paragraph("")

    # ─═ V ═─
    add_heading_styled(doc, "V", 0)
    tbl = create_table(doc, ["Término", "Definición"])
    add_term(tbl, "venv (virtual environment)",
        "Entorno virtual de Python que aísla las dependencias del proyecto. "
        "Se crea con python3 -m venv venv y se activa con source venv/bin/activate.")
    add_term(tbl, "VERDICT (veredicto)",
        "Resultado de la clasificación. Puede ser: REAL, FALSO, SÁTIRA, ESTAFA, "
        "NO VERIFICABLE (de Groq) o DUDOSO (ajuste del frontend).")
    add_term(tbl, "VerdictDisplay (componente)",
        "Componente React que muestra el veredicto con estilos según el tipo: "
        "color, borde, boxShadow, clipPath y textShadow dinámicos. "
        "Usa un diccionario de configuración (VERDICT_CONFIG) por tipo de veredicto.")
    add_term(tbl, "VERIFEX",
        "Nombre del sistema: Analizador de Credibilidad de Noticias. "
        "Backend Python/Flask + Frontend React/TypeScript. Clasifica noticias usando "
        "IA (Groq API) con scraping multicapa y verificación cruzada.")
    add_term(tbl, "Vite",
        "Build tool para frontend. VERIFEX lo usa con plugin de React. "
        "Configura proxy de desarrollo (/analyze → localhost:5001), hot reload, "
        "y build de producción optimizado.")
    add_term(tbl, "VITE_API_URL",
        "Variable de entorno para producción que define la URL del backend. "
        "El frontend la usa para apuntar a https://backend.onrender.com en lugar "
        "de localhost:5001. Obligatoria en producción.")
    add_term(tbl, "Vitest",
        "Framework de pruebas para Vite. VERIFEX tiene 52 tests de frontend "
        "configurados con jsdom, testing-library y setup automático.")
    doc.add_paragraph("")

    # ─═ W ═─
    add_heading_styled(doc, "W", 0)
    tbl = create_table(doc, ["Término", "Definición"])
    add_term(tbl, "web scraping",
        "Técnica de extracción automatizada de datos de sitios web. VERIFEX implementa "
        "5 estrategias (cloudscraper, curl_cffi, requests, Playwright, Google Cache) "
        "para garantizar la extracción exitosa del contenido de la URL analizada.")
    add_term(tbl, "Web Service (Render)",
        "Tipo de servicio en Render para aplicaciones backend. VERIFEX lo usa para "
        "desplegar Flask + Gunicorn. Render ejecuta build.sh y luego el comando del Procfile.")
    doc.add_paragraph("")

    # ─═ X ═─
    add_heading_styled(doc, "X", 0)
    tbl = create_table(doc, ["Término", "Definición"])
    add_term(tbl, "X / Twitter",
        "Red social cuyo dominio (x.com, twitter.com) está en SOCIAL_MEDIA_DOMAINS. "
        "Recibe el prompt especial SOCIAL_MEDIA_PROMPT con reglas de clasificación "
        "específicas para contenido de la plataforma.")
    add_term(tbl, "XML (eXtensible Markup Language)",
        "Formato de marcado usado por RSS. VERIFEX parsea XML de Google News RSS "
        "con xml.etree.ElementTree para extraer noticias relacionadas.")
    doc.add_paragraph("")

    # ─═ Z ═─
    add_heading_styled(doc, "Z", 0)
    tbl = create_table(doc, ["Término", "Definición"])
    add_term(tbl, "zero downtime",
        "Estrategia de despliegue donde la aplicación no deja de estar disponible "
        "durante actualizaciones. Render maneja esto automáticamente con balanceo "
        "entre instancias viejas y nuevas durante el deploy.")
    doc.add_paragraph("")

    filepath = os.path.join(OUT, "Glosario_Terminos_VERIFEX.docx")
    doc.save(filepath)
    print(f"Documento generado: {filepath}")

if __name__ == "__main__":
    main()
