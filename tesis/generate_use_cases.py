#!/usr/bin/env python3
"""Genera Especificaciones de Casos de Uso (formato .docx) - VERIFEX"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUT = os.path.dirname(os.path.abspath(__file__))

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── Portada ──
for _ in range(6):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('VERIFEX')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0, 51, 102)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Analizador de Credibilidad de Noticias')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(100, 100, 100)
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Especificación de Casos de Uso')
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0, 51, 102)
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Metodología: Kanban')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)
for _ in range(4):
    doc.add_paragraph()
for label, value in [('Versión:', '1.0'), ('Fecha:', 'Julio 2026'),
                     ('Equipo:', 'Marco, Luis, Ulises, Tony'),
                     ('Universidad:', 'Universidad Tres Culturas (UTC)')]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'{label} ')
    r.bold = True; r.font.size = Pt(11)
    r = p.add_run(value)
    r.font.size = Pt(11)
doc.add_page_break()

# ── Helper ──
def add_use_case(uc_id, name, actor, desc, trigger, preconds, postconds, main_flow, alt_flows, biz_rules):
    num = int(uc_id.split("-")[1])
    doc.add_heading(f'{uc_id}: {name}', level=1)
    for label, value in [('ID', uc_id), ('Nombre', name), ('Actor', actor),
                         ('Descripción', desc), ('Disparador', trigger),
                         ('Precondiciones', preconds), ('Postcondiciones', postconds)]:
        p = doc.add_paragraph()
        r = p.add_run(f'{label}: ')
        r.bold = True
        if isinstance(value, list):
            for v in value:
                doc.add_paragraph(v, style='List Bullet')
        else:
            p.add_run(value)
    doc.add_paragraph()
    p = doc.add_paragraph(); r = p.add_run('Flujo Principal'); r.bold = True
    for step in main_flow:
        doc.add_paragraph(step, style='List Number')
    if alt_flows:
        doc.add_paragraph()
        p = doc.add_paragraph(); r = p.add_run('Flujos Alternativos'); r.bold = True
        for af_title, af_steps in alt_flows:
            p = doc.add_paragraph(); r = p.add_run(af_title); r.bold = True
            for step in af_steps:
                doc.add_paragraph(step, style='List Bullet')
    if biz_rules:
        doc.add_paragraph()
        p = doc.add_paragraph(); r = p.add_run('Reglas de Negocio'); r.bold = True
        for rule in biz_rules:
            doc.add_paragraph(rule, style='List Bullet')

# ── CU-01 ──
add_use_case('CU-01', 'Analizar URL', 'Usuario',
    'El usuario ingresa una URL de una noticia y el sistema la procesa para determinar su credibilidad.',
    'El usuario hace clic en el botón "Analizar" o presiona Enter después de pegar una URL.',
    ['La URL debe comenzar con http:// o https://.', 'Conexión a internet.',
     'GROQ_API_KEY configurada en el servidor.'],
    ['El sistema muestra veredicto, nivel de confianza, detalles y noticias similares.'],
    ['1. El usuario ingresa una URL en el campo de texto.',
     '2. El usuario hace clic en "Analizar" (o presiona Enter).',
     '3. El sistema valida que la URL comience con http:// o https://.',
     '4. El servidor recibe la solicitud POST /analyze.',
     '5. El sistema ejecuta scrape_url() para obtener el contenido de la URL.',
     '   5a. Intenta cloudscraper (4 perfiles de navegador).',
     '   5b. Si falla, intenta curl_cffi (4 versiones de impersonación).',
     '   5c. Si falla, intenta requests estándar.',
     '   5d. Si falla, intenta Playwright (Firefox headless).',
     '6. El sistema extrae el contenido HTML con BeautifulSoup.',
     '7. El sistema ejecuta find_similar_news() contra Google News RSS.',
     '8. El sistema construye el prompt para Groq con el contenido extraído.',
     '9. El sistema llama a Groq API (intenta llama-3.3-70b, fallback llama-3.1-8b).',
     '10. El sistema parsea la respuesta JSON de Groq.',
     '11. Si el dominio es confiable y el veredicto es FALSO, lo sobrescribe a NO VERIFICABLE.',
     '12. El sistema devuelve el resultado completo al frontend.',
     '13. El frontend renderiza todos los resultados.'],
    [('A-01: URL inválida', ['Muestra "Ingresa una URL válida que comience con http:// o https://".']),
     ('A-02: Timeout', ['Si el scraping excede 60s, muestra "La URL tardó demasiado en responder".']),
     ('A-03: Error HTTP', ['Captura el error HTTP y muestra un mensaje descriptivo (404, 403, etc.).']),
     ('A-04: Fallan todos los scrapers', ['Muestra los errores de cada scraper.']),
     ('A-05: Contenido insuficiente', ['Si el texto extraído es < 100 caracteres, intenta Playwright.']),
     ('A-06: Error en Groq', ['Si ambos modelos fallan, muestra "Error al analizar con la IA".']),
     ('A-07: Sin conexión', ['Si el fetch falla, muestra "Error de conexión".'])],
    ['RN-01: La URL debe comenzar con http:// o https://.',
     'RN-02: No se almacena ningún dato del usuario.',
     'RN-03: Dominio confiable + FALSO = NO VERIFICABLE.',
     'RN-04: Confianza < 50 fuerza veredicto a FALSO (excepto Instagram).',
     'RN-05: Confianza 50-69 fuerza veredicto a DUDOSO (excepto Instagram).'])

# ── CU-02 ──
add_use_case('CU-02', 'Ver Veredicto de Credibilidad', 'Usuario',
    'El usuario visualiza el veredicto final: REAL, FALSO, SÁTIRA, ESTAFA, DUDOSO o NO VERIFICABLE.',
    'CU-01 completado.',
    ['CU-01 ejecutado correctamente.', 'LLM devolvió un veredicto.'],
    ['El usuario conoce la credibilidad de la noticia.'],
    ['1. El sistema recibe el veredicto del análisis.',
     '2. Si confianza < 50, cambia a FALSO (excepto Instagram).',
     '3. Si confianza 50-69, cambia a DUDOSO (excepto Instagram).',
     '4. Muestra el veredicto final con su color correspondiente.',
     '5. Si se ajustó, muestra el veredicto original en letra pequeña.'],
    [('A-01: Formato inválido', ['Si Groq devuelve JSON mal formado, intenta parse_response() para extraer el JSON.'])],
    ['RN-06: Veredictos posibles: REAL, FALSO, SÁTIRA, ESTAFA, NO VERIFICABLE.',
     'RN-07: Colores: REAL=cian, FALSO=rojo, SÁTIRA=naranja, ESTAFA=morado, NO VERIFICABLE=gris.'])

# ── CU-03 ──
add_use_case('CU-03', 'Visualizar Nivel de Confianza', 'Usuario',
    'El usuario ve un puntaje 0-100 con barra de progreso de 20 segmentos.',
    'CU-01 completado.', ['CU-01 ejecutado correctamente.'],
    ['El usuario conoce la certeza del análisis.'],
    ['1. El frontend recibe confidence_score.',
     '2. Renderiza el número (ej: "87%").',
     '3. Renderiza barra de 20 segmentos coloreados: rojo (0-40), naranja (41-69), cian (70-89), verde (90-100).'],
    [], ['RN-08: Confianza es entero 0-100.', 'RN-09: Barra de 20 segmentos (5 pts c/u).'])

# ── CU-04 ──
add_use_case('CU-04', 'Leer Resumen del Artículo', 'Usuario',
    'El usuario lee el resumen generado por el LLM.',
    'CU-01 completado.', ['CU-01 ejecutado correctamente.'],
    ['El usuario entiende de qué trata el artículo sin leerlo completo.'],
    ['1. El frontend recibe el campo "summary".', '2. Muestra el resumen en una tarjeta de texto.'],
    [('A-01: Resumen vacío', ['Muestra "No disponible".'])],
    ['RN-10: Resumen generado por LLM, ~200 caracteres.'])

# ── CU-05 ──
add_use_case('CU-05', 'Consultar Afirmaciones Principales', 'Usuario',
    'El usuario revisa las afirmaciones clave extraídas del artículo.',
    'CU-01 completado.', ['CU-01 ejecutado correctamente.', 'LLM extrajo afirmaciones.'],
    ['El usuario identifica los puntos clave.'],
    ['1. El frontend recibe "extracted_claims".', '2. Renderiza cada afirmación como viñeta.'],
    [('A-01: Sin afirmaciones', ['Muestra "No se pudieron extraer afirmaciones".'])],
    ['RN-11: Afirmaciones como lista con viñetas.'])

# ── CU-06 ──
add_use_case('CU-06', 'Revisar Análisis Detallado', 'Usuario',
    'El usuario lee el razonamiento paso a paso detrás del veredicto.',
    'CU-01 completado.', ['CU-01 ejecutado correctamente.', 'LLM generó razonamiento.'],
    ['El usuario entiende los criterios del veredicto.'],
    ['1. El frontend recibe "reasoning".', '2. Renderiza cada punto como elemento numerado.'],
    [('A-01: Sin razonamiento', ['Muestra "No disponible".'])],
    ['RN-12: Razonamiento como lista numerada.'])

# ── CU-07 ──
add_use_case('CU-07', 'Identificar Alertas y Señales Positivas', 'Usuario',
    'El usuario visualiza banderas rojas y señales positivas en formato de etiquetas (pills).',
    'CU-01 completado.',
    ['CU-01 ejecutado correctamente.', 'LLM identificó banderas rojas y/o señales positivas.'],
    ['El usuario conoce aspectos problemáticos y confiables.'],
    ['1. El frontend recibe "red_flags" y "positive_signals".',
     '2. Renderiza banderas rojas en panel rojo con pills.',
     '3. Renderiza señales positivas en panel verde con pills.'],
    [('A-01: Sin banderas rojas', ['Muestra "Sin alertas detectadas".']),
     ('A-02: Sin señales positivas', ['Muestra "Sin señales positivas".'])],
    ['RN-13: Banderas rojas en tonos rojos, señales positivas en verdes.'])

# ── CU-08 ──
add_use_case('CU-08', 'Explorar Noticias Similares', 'Usuario',
    'El usuario navega una cuadrícula de noticias relacionadas de Google News RSS.',
    'CU-01 completado.', ['CU-01 ejecutado correctamente.', 'Título extraído correctamente.'],
    ['El usuario tiene contexto adicional de otras fuentes.'],
    ['1. Backend ejecuta find_similar_news(title, max_results=5).',
     '2. Google News RSS devuelve XML con resultados.',
     '3. Frontend renderiza cuadrícula con tarjetas (fuente, título, fecha).',
     '4. Al hacer clic, abre la URL en nueva pestaña.'],
    [('A-01: Sin resultados', ['Muestra "No se encontraron noticias similares".']),
     ('A-02: Error RSS', ['Devuelve arreglo vacío, se omite la sección.'])],
    ['RN-14: Hasta 5 resultados, región MX, español.',
     'RN-15: Carga diferida (lazy load).'])

# ── CU-09 ──
add_use_case('CU-09', 'Cambiar Idioma (ES/EN)', 'Usuario',
    'El usuario alterna entre español e inglés en toda la interfaz.',
    'El usuario hace clic en el botón ES/EN.',
    ['App cargada en el navegador.', 'Botón de idioma accesible.'],
    ['Toda la interfaz en el idioma seleccionado.'],
    ['1. Usuario hace clic en "ES" o "EN".',
     '2. Frontend actualiza estado "lang".',
     '3. Todos los textos se actualizan al instante.',
     '4. El botón refleja el idioma actual.'],
    [], ['RN-16: Traducciones en objeto dentro de App.tsx.',
         'RN-17: Cambio instantáneo sin recargar.', 'RN-18: Valor predeterminado: español.'])

# ── CU-10 ──
add_use_case('CU-10', 'Consultar Estado del Sistema', 'Usuario/Sistema',
    'Verifica que el servidor backend está funcionando mediante GET /health.',
    'Solicitud GET a /health.', ['Servidor Flask corriendo.'],
    ['Se confirma que el servidor está operativo.'],
    ['1. Cliente envía GET a /health.',
     '2. Servidor responde {"status": "ok"} con código 200.',
     '3. Cliente confirma servidor activo.'],
    [('A-01: Servidor caído', ['No hay respuesta, error de conexión.'])],
    ['RN-19: /health no requiere autenticación.',
     'RN-20: Siempre responde {"status": "ok"} con 200.'])

# ── Guardar ──
output = os.path.join(OUT, 'Especificacion_Casos_de_Uso_VERIFEX.docx')
doc.save(output)
print(f'✅ {output}')
