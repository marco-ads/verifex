from docx import Document
from docx.shared import Pt, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree
import re

SRC = '/Users/maarco_serrano/Downloads/verifex-standalone 2/tesis/Tesis-Ulikes.docx'
DST = '/Users/maarco_serrano/Downloads/verifex-standalone 2/tesis/Tesis-Ulikes.docx'

doc = Document(SRC)

FONT_NAME = 'Times New Roman'
FONT_SIZE = Pt(12)

# ========================================
# 1. PAGE SETUP — 1 inch margins
# ========================================
for sec in doc.sections:
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)

# ========================================
# 2. STYLE CONFIGURATION
# ========================================

def cfg_style(style, bold=None, italic=None, align=None, first_indent=None,
              left_indent=None, line_spacing=2.0, space_before=0, space_after=0):
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic
    pf = style.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if align is not None:
        pf.alignment = align
    if first_indent is not None:
        pf.first_line_indent = first_indent
    if left_indent is not None:
        pf.left_indent = left_indent

# Normal — base style
cfg_style(doc.styles['Normal'], align=WD_ALIGN_PARAGRAPH.LEFT)

# Heading 1 → APA Level 1: Centered, Bold
cfg_style(doc.styles['Heading 1'], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
          first_indent=Inches(0))
# Heading 2 → APA Level 2: Left, Bold
cfg_style(doc.styles['Heading 2'], bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
          first_indent=Inches(0))
# Heading 3 → APA Level 3: Left, Bold, Italic
cfg_style(doc.styles['Heading 3'], bold=True, italic=True, align=WD_ALIGN_PARAGRAPH.LEFT,
          first_indent=Inches(0))

# Body Text — with first-line indent for body paragraphs
cfg_style(doc.styles['Body Text'], align=WD_ALIGN_PARAGRAPH.LEFT,
          first_indent=Inches(0.5))

# List Bullet — no first-line indent
cfg_style(doc.styles['List Bullet'], align=WD_ALIGN_PARAGRAPH.LEFT,
          first_indent=Inches(0))

# List Paragraph
cfg_style(doc.styles['List Paragraph'], align=WD_ALIGN_PARAGRAPH.LEFT)

# TOC styles
for sname in ['toc 1', 'toc 2', 'toc 3', 'TOC Heading']:
    try:
        cfg_style(doc.styles[sname], align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=Inches(0))
    except KeyError:
        pass

# ========================================
# 3. HELPER FUNCTIONS
# ========================================

def para_in_table(para):
    p = para._p
    while p is not None:
        if p.tag.endswith('}tc'):
            return True
        p = p.getparent()
    return False

def para_has_image(para):
    for child in para._p.iter():
        if child.tag in [qn('w:drawing'), qn('w:pict'), qn('wp:inline'), qn('wp:anchor')]:
            return True
    return False

def para_has_only_tab_or_space(text):
    return bool(re.match(r'^[\s\t]+$', text))

def is_heading_para(para):
    return para.style and para.style.name.startswith('Heading')

CHAPTER_RE = re.compile(r'^(Capítulo|CAPÍTULO)\s+\d', re.IGNORECASE)
SECTION_TITLE_NAMES = {'agradecimiento', 'dedicatoria', 'resumen', 'abstract',
                        'índice', 'indice', 'tabla de contenido',
                        'referencias', 'bibliografía', 'bibliografia',
                        'apéndice', 'apendice', 'anexo', 'conclusiones',
                        'introducción', 'introduccion', 'glosario',
                        'manual técnico', 'manual tecnico', 'índice de tablas',
                        'índice de figuras', 'indice de tablas', 'indice de figuras'}

def is_section_title(text):
    t = text.strip().lower().rstrip('.')
    return t in SECTION_TITLE_NAMES

def is_chapter_title(text):
    return bool(CHAPTER_RE.match(text.strip()))

# ========================================
# 4. PARAGRAPH-BY-PARAGRAPH PROCESSING
# ========================================
#
# Rules:
#   - Do NOT touch paragraphs inside tables
#   - Do NOT touch paragraphs containing images
#   - Heading paragraphs: their style handles everything
#   - Section titles (Agradecimiento, Referencias, etc.): set to centered, bold
#   - Chapter titles (Capítulo N: ...): set to Heading 1 style
#   - All other body text: left-align, TNR 12pt, double-space, 0.5" first-line indent
#   - List items: no first-line indent
#   - JUSTIFY → LEFT everywhere
#

for i, para in enumerate(doc.paragraphs):
    text_raw = para.text
    text = text_raw.strip()

    # === SKIP conditions ===
    if para_in_table(para) or para_has_image(para):
        continue

    style_name = para.style.name if para.style else ''
    is_heading = style_name.startswith('Heading')
    is_list = style_name in ('List Bullet', 'List Paragraph')
    is_normal = style_name == 'Normal'
    is_body = style_name == 'Body Text'

    # === FIX JUSTIFY → LEFT (only if not centered) ===
    if para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Also fix via runs if needed
    for run in para.runs:
        rPr = run._r.find(qn('w:rPr'))
        if rPr is not None:
            jc = rPr.find(qn('w:jc'))
            if jc is not None:
                rPr.remove(jc)

    # === FONT & SIZE: ensure TNR 12pt on all runs ===
    for run in para.runs:
        rPr = run._r.find(qn('w:rPr'))
        if rPr is not None:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                for attr in ['ascii', 'hAnsi', 'cs', 'eastAsia']:
                    val = rFonts.get(qn(f'w:{attr}'))
                    if val and val not in (FONT_NAME, 'Symbol', 'Courier New', 'Lucida Console', 'Consolas'):
                        rFonts.set(qn(f'w:{attr}'), FONT_NAME)
            sz = rPr.find(qn('w:sz'))
            if sz is not None:
                val = sz.get(qn('w:val'))
                if val and int(val) > 26:
                    sz.set(qn('w:val'), '24')
                # val 24 = 12pt (half-points)

    # === SPECIAL: Section titles (Agradecimiento, Referencias, etc.) ===
    if is_section_title(text) and not is_heading:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Set bold on all runs
        for run in para.runs:
            run.bold = True
        continue

    # === SPECIAL: Chapter titles "Capítulo N: ..." without Heading style ===
    if is_chapter_title(text) and not is_heading:
        para.style = doc.styles['Heading 1']
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        continue

    # === HEADINGS: skip (style handles formatting) ===
    if is_heading:
        continue

    # === LISTS: skip (style handles) ===
    if is_list:
        continue

    # === EMPTY or whitespace-only: skip ===
    if not text or para_has_only_tab_or_space(text_raw):
        continue

    # === BODY TEXT: apply left-align, double-spacing, first-line indent ===
    if is_body or is_normal:
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 2.0

        if para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            para.paragraph_format.first_line_indent = Inches(0.5)

# ========================================
# 5. PAGE NUMBERS IN HEADER (upper right)
# ========================================
# APA 7 requires page number in upper right corner of every page
for sec in doc.sections:
    header = sec.header
    header.is_linked_to_previous = False
    # Clear existing header paragraphs
    for p in header.paragraphs:
        p.clear()
    # Add page number field
    if header.paragraphs:
        p = header.paragraphs[0]
    else:
        p = header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    run = p.add_run()
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE

    fldChar_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar_begin)

    run2 = p.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instrText)

    run3 = p.add_run()
    fldChar_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fldChar_end)

# ========================================
# 6. SAVE
# ========================================
doc.save(DST)
print("APA 7 formatting applied successfully.")
