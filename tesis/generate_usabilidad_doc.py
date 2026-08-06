#!/usr/bin/env python3
"""Genera documento .docx con el Estudio de Usabilidad de VERIFEX."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.dirname(os.path.abspath(__file__))

NAVY = "1B3A5C"
DARK_SLATE = "2C3E50"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADING_COLOR = RGBColor(0x1B, 0x3A, 0x5C)
SUBTITLE_GRAY = RGBColor(0x70, 0x70, 0x80)

def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)

def add_attr_row(table, *cell_texts):
    row = table.add_row()
    for i, text in enumerate(cell_texts):
        if i < len(row.cells):
            row.cells[i].text = text
    for cell in row.cells:
        for p in cell.paragraphs:
            p.style.font.size = Pt(8.5)
            p.style.font.name = "Consolas"

def create_table(doc, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = 1
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = "Consolas"
            run.font.color.rgb = WHITE
        set_cell_shading(cell, DARK_SLATE)
    return table

def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = HEADING_COLOR
    return h

def main():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    # ── Portada ──
    for _ in range(6):
        doc.add_paragraph("")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("VERIFEX")
    run.font.size = Pt(42)
    run.bold = True
    run.font.color.rgb = HEADING_COLOR

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Estudio de Usabilidad")
    run.font.size = Pt(18)
    run.font.color.rgb = SUBTITLE_GRAY

    doc.add_paragraph("")
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub2.add_run("Herramientas de acuerdo a las necesidades del levantamiento")
    run.font.size = Pt(14)
    run.font.color.rgb = SUBTITLE_GRAY

    doc.add_paragraph("")
    sub3 = doc.add_paragraph()
    sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub3.add_run("Backend Python/Flask + Frontend React/TypeScript")
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = SUBTITLE_GRAY

    doc.add_page_break()

    # ── 1. Introducción ──
    add_heading_styled(doc, "1. Introducción", 0)
    doc.add_paragraph(
        "El presente estudio de usabilidad describe las herramientas tecnológicas seleccionadas "
        "para cada módulo del sistema VERIFEX, justificando su elección con base en los requisitos "
        "funcionales y no funcionales identificados durante el levantamiento de requerimientos. "
        "Cada herramienta fue evaluada según criterios de idoneidad, rendimiento, curva de "
        "aprendizaje, comunidad y costo."
    )
    doc.add_paragraph("")

    # ── 2. Requerimientos del Levantamiento ──
    add_heading_styled(doc, "2. Requerimientos del Levantamiento", 0)
    doc.add_paragraph(
        "El levantamiento de requerimientos identificó las siguientes necesidades clave que "
        "determinaron la selección de herramientas:"
    )
    doc.add_paragraph("")
    tbl = create_table(doc, ["Necesidad", "Requerimiento asociado", "Herramienta seleccionada", "Justificación"])
    add_attr_row(tbl,
        "Analizar credibilidad de noticias desde una URL",
        "HU-01, HU-02, HU-18",
        "Flask + Playwright + cloudscraper + curl_cffi + requests",
        "Se requiere un backend ligero que procese URLs robustamente. "
        "La pila de scraping multicapa garantiza extracción exitosa incluso con "
        "protecciones Cloudflare, bloqueos por IP o JavaScript pesado.")
    add_attr_row(tbl,
        "Clasificar contenido con IA sin GPU local",
        "HU-03, HU-19, HU-20",
        "Groq API (llama-3.3-70b-versatile + fallback llama-3.1-8b-instant)",
        "Groq ofrece inferencia ultrarrápida en la nube sin necesidad de hardware "
        "especializado. El fallback garantiza disponibilidad. Temperature baja (0.1) y "
        "response_format='json_object' aseguran respuestas consistentes y parseables.")
    add_attr_row(tbl,
        "Interfaz intuitiva y responsive",
        "HU-11, HU-12, HU-15, HU-27, HU-28, HU-29",
        "React 18 + TypeScript + Vite + Tailwind CSS",
        "React permite componentes reutilizables con estado local. TypeScript añade "
        "type safety. Vite optimiza builds y recarga en caliente. Tailwind agiliza el "
        "diseño responsive con utilidades directas en JSX.")
    add_attr_row(tbl,
        "Búsqueda de contexto adicional",
        "HU-08, HU-22",
        "Google News RSS (feedparser)",
        "API gratuita, sin autenticación, devuelve resultados en español de México "
        "(hl=es, gl=MX). Integración directa vía feedparser sin dependencias pesadas.")
    add_attr_row(tbl,
        "Comunicación frontend-backend segura",
        "HU-16, HU-17, HU-24, HU-25",
        "Fetch API + AbortController + Flask-CORS",
        "Fetch nativo del navegador evita librerías extras. AbortController maneja "
        "timeouts de 60s. Flask-CORS permite orígenes cruzados en desarrollo y producción.")
    add_attr_row(tbl,
        "Despliegue en la nube sin administración de servidores",
        "HU-23, DEP-01, DEP-02, DEP-03",
        "Render (Web Service + Static Site) + Gunicorn",
        "Render ofrece despliegue gratuito con build automático, SSL y dominio público. "
        "Gunicorn sirve Flask con múltiples workers para producción.")
    add_attr_row(tbl,
        "Pruebas automatizadas",
        "HU-34, TEST-01, TEST-02, TEST-03",
        "pytest + Vitest + Testing Library",
        "pytest para backend (27 tests), Vitest + Testing Library para frontend "
        "(52 tests en 7 archivos). Cobertura de integración completa.")
    doc.add_paragraph("")

    # ── 3. Criterios de Selección ──
    add_heading_styled(doc, "3. Criterios de Selección", 0)
    doc.add_paragraph(
        "Cada herramienta fue evaluada bajo los siguientes criterios antes de ser incorporada al proyecto:"
    )
    doc.add_paragraph("")
    tbl = create_table(doc, ["Criterio", "Descripción", "Peso"])
    add_attr_row(tbl, "Idoneidad funcional",
        "La herramienta resuelve directamente la necesidad identificada en el levantamiento. "
        "Se priorizaron herramientas con propósito específico sobre soluciones genéricas.",
        "Alto")
    add_attr_row(tbl, "Rendimiento",
        "Tiempo de respuesta, consumo de recursos y escalabilidad. "
        "Se midió latencia de Groq API (~2-5s por análisis), tiempo de scraping "
        "(~10-30s con fallbacks) y tamaño de build de producción (~200KB JS).",
        "Alto")
    add_attr_row(tbl, "Madurez y comunidad",
        "Versión estable, documentación actualizada, número de usuarios y "
        "frecuencia de actualizaciones. Todas las herramientas seleccionadas "
        "tienen más de 5 años de desarrollo activo.",
        "Medio")
    add_attr_row(tbl, "Costo",
        "Las herramientas seleccionadas son gratuitas (open source) o tienen "
        "tier gratuitos generosos. Groq API ofrece crédito gratuito inicial; "
        "Render tiene tier free con limitaciones razonables.",
        "Medio")
    add_attr_row(tbl, "Curva de aprendizaje",
        "Tiempo estimado para que un desarrollador con experiencia web básica "
        "sea productivo. Flask y React son ampliamente conocidos; Playwright "
        "requiere configuración adicional (instalación de navegadores).",
        "Bajo")
    doc.add_paragraph("")

    # ── 4. Alternativas Evaluadas y Descartadas ──
    add_heading_styled(doc, "4. Alternativas Evaluadas y Descartadas", 0)
    doc.add_paragraph(
        "Durante el levantamiento se evaluaron las siguientes alternativas que fueron descartadas "
        "por no cumplir con los criterios de selección:"
    )
    doc.add_paragraph("")
    tbl = create_table(doc, ["Herramienta", "Módulo", "Motivo de descarte"])
    add_attr_row(tbl, "Django + Django REST Framework", "Backend",
        "Sobredimensionado para un API con solo 2 endpoints. Flask es más ligero, "
        "tiene menor curva de aprendizaje y arranca en segundos.")
    add_attr_row(tbl, "Next.js / Nuxt.js", "Frontend",
        "Renderizado SSR innecesario para una SPA de una sola vista. "
        "React + Vite es más simple y el build estático es suficiente.")
    add_attr_row(tbl, "Selenium", "Scraping",
        "Más pesado que Playwright, requiere WebDriver y es más lento. "
        "Playwright tiene mejor API para esperar condiciones y manejar Cloudflare.")
    add_attr_row(tbl, "OpenAI API / Claude API", "Clasificación IA",
        "Mayor latencia y costo por consulta. Groq ofrece mejor rendimiento "
        "para tareas de clasificación estructurada con respuesta JSON.")
    add_attr_row(tbl, "Heroku / Railway", "Despliegue",
        "Heroku ya no tiene tier gratuito. Railway tiene limitaciones de memoria "
        "que causaron fallos en el despliegue del scraper con Playwright.")
    add_attr_row(tbl, "Axios", "HTTP frontend",
        "Fetch API nativo es suficiente para una sola llamada POST. "
        "Axios añadiría ~14KB innecesarios al bundle.")
    doc.add_paragraph("")

    # ── 5. Mapeo de Herramientas a Requerimientos Funcionales ──
    add_heading_styled(doc, "5. Mapeo de Herramientas a Requerimientos Funcionales", 0)
    doc.add_paragraph(
        "La siguiente tabla resume la correspondencia entre los requerimientos funcionales "
        "(HU) y las herramientas que los implementan:"
    )
    doc.add_paragraph("")
    tbl = create_table(doc, ["HU", "Descripción", "Herramienta(s)", "Módulo"])
    mappings = [
        ("HU-01, HU-02", "Pegar URL y ver veredicto", "React + Flask + Groq", "Frontend/Backend"),
        ("HU-03", "Nivel de confianza (20 niveles)", "React (ConfidenceBar)", "Frontend"),
        ("HU-04", "Banderas rojas y señales positivas", "React (RedFlags) + Groq", "Frontend/Backend"),
        ("HU-05, HU-06", "Resumen y afirmaciones con IA", "Groq API + analyzer.py", "Backend"),
        ("HU-07", "Análisis por categoría", "Groq API (prompt engineering)", "Backend"),
        ("HU-08, HU-22", "Noticias similares", "Google News RSS + news_finder.py", "Backend"),
        ("HU-09", "Tipo de artículo", "Groq API (clasificación)", "Backend"),
        ("HU-10", "Alerta de estafa", "Groq API (is_scam) + React", "Frontend/Backend"),
        ("HU-11", "Cambio de idioma", "React (LanguageToggle)", "Frontend"),
        ("HU-12, HU-26, HU-27, HU-28, HU-29", "Interfaz cyberpunk responsive", "React + Tailwind + CSS custom", "Frontend"),
        ("HU-16", "Sin almacenamiento de datos", "Flask (sin BD) + Fetch API", "Arquitectura"),
        ("HU-17", "Conexión segura HTTPS", "Render (SSL automático)", "Despliegue"),
        ("HU-18", "Scraper robusto", "cloudscraper + curl_cffi + requests + Playwright + Google Cache", "Backend"),
        ("HU-19, HU-20", "Clasificación con IA y prompt engineering", "Groq API + analyzer.py", "Backend"),
        ("HU-21", "Override de dominios confiables", "CREDIBLE_DOMAINS + analyzer.py", "Backend"),
        ("HU-23", "Despliegue zero downtime", "Render + Gunicorn", "Despliegue"),
        ("HU-24", "Backend con /analyze y /health", "Flask + app.py", "Backend"),
        ("HU-25", "Frontend con fetch + timeout", "React + Fetch API + AbortController", "Frontend"),
        ("HU-30 a HU-34", "Documentación, diagramas, pruebas", "pytest + Vitest + UML + Markdown", "Documentación"),
    ]
    for m in mappings:
        add_attr_row(tbl, *m)
    doc.add_paragraph("")

    # ── 6. Conclusión ──
    add_heading_styled(doc, "6. Conclusión", 0)
    doc.add_paragraph(
        "Las herramientas seleccionadas cumplen con los requerimientos del levantamiento "
        "al priorizar simplicidad, rendimiento y costo cero. La arquitectura resultante "
        "es mantenible, extensible y desplegable sin infraestructura compleja. "
        "Las alternativas descartadas fueron evaluadas objetivamente y no se descarta "
        "su incorporación futura si los requerimientos del proyecto evolucionan."
    )

    filepath = os.path.join(OUT, "Estudio_Usabilidad_Herramientas_VERIFEX.docx")
    doc.save(filepath)
    print(f"Documento generado: {filepath}")

if __name__ == "__main__":
    main()
